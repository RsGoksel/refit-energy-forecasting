"""
Generate the comprehensive Final Project Report (DOCX) and convert to PDF.

Clean style:
 - Calibri throughout
 - 1.3 line spacing
 - No em-dash, no superscript R-squared, no et al., no subscript digits
 - "and others" instead of "et al."

Rich content:
 - Cover page
 - Abstract
 - Section structure: Intro -> Data -> Methodology -> Results -> Discussion -> Conclusion
 - All 6 EDA/result figures + 4 custom diagrams
 - Results tables
 - References
"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.abspath(__file__))
RESULTS = os.path.join(BASE, "results")
OUTPUT_DOCX = os.path.join(BASE, "EBT629E_Final_Project_Report.docx")
OUTPUT_PDF = os.path.join(BASE, "EBT629E_Final_Project_Report.pdf")

doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

# Default style
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.line_spacing = 1.3
style.paragraph_format.space_after = Pt(5)

NAVY = RGBColor(0, 51, 102)
GREY = RGBColor(80, 80, 80)


# ============================================================
# HELPER FUNCTIONS
# ============================================================
def add_heading(text, level=1, color=NAVY):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
        run.font.name = 'Calibri'
    return h


def add_para(text, lead_bold=None, justify=True):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if lead_bold:
        r = p.add_run(lead_bold)
        r.bold = True
        r.font.name = 'Calibri'
        r.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.name = 'Calibri'
    r2.font.size = Pt(11)
    return p


def add_figure(filename, caption, width=Inches(6.0)):
    path = os.path.join(RESULTS, filename)
    if not os.path.exists(path):
        doc.add_paragraph(f"[Figure missing: {filename}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run()
    run.add_picture(path, width=width)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    r = cap.add_run(caption)
    r.bold = True
    r.font.size = Pt(9)
    r.font.name = 'Calibri'


def add_bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    for r in p.runs:
        r.font.name = 'Calibri'
        r.font.size = Pt(11)


def add_table_caption(text):
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(2)
    r = cap.add_run(text)
    r.bold = True
    r.font.size = Pt(9)
    r.font.name = 'Calibri'


# ============================================================
# COVER PAGE
# ============================================================
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Istanbul Technical University")
r.font.size = Pt(14)
r.font.name = 'Calibri'
r.font.color.rgb = GREY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("EBT 629E, Artificial Intelligence")
r.font.size = Pt(14)
r.font.name = 'Calibri'
r.font.color.rgb = GREY

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Final Project Report")
r.font.size = Pt(22)
r.bold = True
r.font.name = 'Calibri'
r.font.color.rgb = NAVY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Household Energy Consumption Forecasting\nUsing Time Series Models")
r.font.size = Pt(16)
r.bold = True
r.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("A Comparative Study of ARIMA, LSTM, and FB-Prophet on the REFIT Dataset")
r.font.size = Pt(12)
r.italic = True
r.font.name = 'Calibri'
r.font.color.rgb = GREY

for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Group Members")
r.font.size = Pt(12)
r.bold = True
r.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    "Nurullah Yıldırım (301252004)\n"
    "Kadir Göksel Gündüz (301241077)\n"
    "Furkan Çınar (301212001)"
)
r.font.size = Pt(12)
r.font.name = 'Calibri'

for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("May 2026")
r.font.size = Pt(12)
r.font.name = 'Calibri'
r.font.color.rgb = GREY

doc.add_page_break()

# ============================================================
# ABSTRACT
# ============================================================
add_heading("Abstract", level=1)
add_para(
    "This report presents a comparative study of four time series forecasting approaches, "
    "applied to household energy consumption data from the REFIT Smart Home dataset. We use 22 "
    "months of 8-second resolution measurements from three UK households (Houses 1, 2, and 5), "
    "aggregate them to daily mean power and apply a 7-day rolling smoother, then split each "
    "series temporally into 80 percent training and 20 percent testing. Alongside the three "
    "candidate models, ARIMA, Long Short-Term Memory (LSTM), and FB-Prophet, we include a "
    "persistence baseline (y_hat[t] = y[t-1]) as a sanity check. ARIMA is tuned through an ADF-"
    "based stationarity test and an AIC grid search, then run under a rolling one-step-ahead "
    "protocol; LSTM uses a three-layer stacked recurrent architecture; Prophet uses weekly and "
    "yearly seasonality in single multi-step mode. On House 2, ARIMA(2, 1, 2) reaches R2 = 0.876 "
    "and RMSE = 48.7 W, but a Diebold-Mariano test shows its edge over the persistence baseline "
    "(R2 = 0.873) is not statistically significant (p = 0.665). LSTM (R2 = 0.748) and Prophet "
    "(R2 = 0.015) lose significantly to persistence (p < 0.001 in both cases). The same pattern "
    "holds on House 1 and House 5: ARIMA narrowly wins, but adds little over persistence, while "
    "LSTM and Prophet underperform the baseline. The headline lesson is that smoothed daily "
    "residential consumption is dominated by lag-1 autocorrelation, so a one-line rule captures "
    "almost all of the predictable structure. Heavier machinery is justified only when stronger "
    "exogenous inputs or finer-grained targets break that autocorrelation. The full code base, "
    "with fixed random seed and reproducible pipeline, is provided alongside this report."
)

doc.add_page_break()

# ============================================================
# 1. INTRODUCTION
# ============================================================
add_heading("1. Introduction", level=1)

add_heading("1.1 Background and Motivation", level=2)
add_para(
    "Household electricity consumption sits at the centre of several modern energy challenges. "
    "On the supply side, accurate short-term forecasts are needed for grid operators to balance "
    "renewable generation, plan reserve capacity, and reduce curtailment. On the demand side, "
    "they enable consumers to participate in demand-response programmes, schedule appliances "
    "intelligently, and optimise integration with rooftop photovoltaic and battery storage "
    "systems. The widespread roll-out of smart meters across Europe and the increasing maturity "
    "of machine learning frameworks have made it feasible to study this problem at the level of "
    "individual households rather than at the substation level."
)
add_para(
    "Forecasting residential consumption is, however, structurally harder than forecasting "
    "aggregated demand. A single household exhibits abrupt switching events (kettles, ovens, "
    "washing machines), idiosyncratic schedules driven by occupant behaviour, and seasonality "
    "that interacts with weather and holidays. Smoothing across many households averages much "
    "of this noise away, which is one reason that grid-level forecasts often report low MAPE "
    "while household-level forecasts can fail outright."
)

add_heading("1.2 Project Objectives", level=2)
add_para(
    "The present project asks a focused, practical question: among three widely used time series "
    "models, ARIMA, LSTM, and FB-Prophet, which is best suited to forecasting the daily energy "
    "consumption of a single household, and why? Concretely, the project pursues four objectives."
)
add_bullet("Build a reproducible pipeline from raw 8-second REFIT readings to model outputs.")
add_bullet("Compare ARIMA, LSTM, and FB-Prophet on a single test bed under the same data split.")
add_bullet("Quantify the gap between models using MSE, RMSE, MAE, and R2 metrics.")
add_bullet("Investigate where each model wins or loses, and document the iterations that led to "
           "the final design.")

add_heading("1.3 Report Structure", level=2)
add_para(
    "Section 2 introduces the REFIT dataset and the chosen household. Section 3 details the "
    "preprocessing pipeline and the three modelling approaches, including the rolling forecast "
    "protocol that distinguishes our setup from many of the multi-step benchmarks reported in "
    "the literature. Section 4 presents the exploratory data analysis. Section 5 reports the "
    "quantitative results and visual comparisons. Section 6 discusses the findings, the "
    "limitations of the study, and threats to validity. Section 7 closes with conclusions and "
    "directions for further work."
)

# ============================================================
# 2. DATASET
# ============================================================
add_heading("2. Dataset: REFIT Smart Home Electrical Load Measurements", level=1)

add_heading("2.1 Overview", level=2)
add_para(
    "REFIT (Personalised Retrofit Decision Support Tools for UK Homes) is a longitudinal dataset "
    "collected between October 2013 and June 2015 in the Loughborough area of the United Kingdom. "
    "The data was produced by a consortium of the Universities of Strathclyde, Loughborough, and "
    "East Anglia under EPSRC funding and is released under a Creative Commons Attribution 4.0 "
    "International licence. Twenty households were instrumented (numbered House 1 to House 21, "
    "skipping House 14) with one current clamp on the whole-house feed and nine Individual "
    "Appliance Monitors per home. Active power was sampled in Watts at roughly one reading every "
    "six to eight seconds, which produced a corpus of about 1.19 billion observations across all "
    "houses (Murray, Stankovic, and Stankovic, 2017)."
)

add_heading("2.2 Selected Household: House 2", level=2)
add_para(
    "We use House 2 as the test bed for this project. The CSV file is 299 MB on disk and contains "
    "5,733,526 timestamped readings spanning 22 months. The monitored appliances are Fridge-"
    "Freezer, Washing Machine, Dishwasher, Television, Microwave, Toaster, Hi-Fi, Kettle, and "
    "Oven Extractor Fan. House 2 was selected because it has no documented signature changes, no "
    "rooftop PV interference, and represents a typical UK household appliance mix."
)

# Correlation heatmap
add_figure("correlation_heatmap.png",
           "Figure 1. Pearson correlation among the aggregate power channel and the nine "
           "individual appliance monitors in House 2.",
           width=Inches(5.5))

add_heading("2.3 Data Quality and Cleaning", level=2)
add_para(
    "The REFIT release applies forward-filling for short gaps and zeros out gaps longer than "
    "two minutes. Readings above 4000 W on appliance channels (above sensor range) have already "
    "been removed by the original maintainers. Our additional cleaning steps remove negative "
    "aggregate readings (caused by clamp polarity glitches) and trim the upper one percent of "
    "values as outliers. After these steps, the aggregate series is continuous and ready for "
    "resampling."
)

doc.add_page_break()

# ============================================================
# 3. METHODOLOGY
# ============================================================
add_heading("3. Methodology", level=1)

add_figure("pipeline_diagram.png",
           "Figure 2. End-to-end machine learning pipeline. Raw 8-second REFIT data flows "
           "through outlier removal, daily resampling, and 7-day smoothing, then feeds three "
           "competing models whose forecasts are evaluated against the held-out test set.",
           width=Inches(6.3))

add_heading("3.1 Data Preprocessing", level=2)
add_para(
    "Five preprocessing steps prepare the raw signal for modelling. First, the timestamp column "
    "is parsed into a pandas DatetimeIndex. Second, the Unix epoch column is dropped to save "
    "memory. Third, outlier filtering removes negative values and entries above the 99th "
    "percentile. Fourth, the cleaned series is resampled to daily means, reducing 5.7 million "
    "rows to roughly 630 daily observations and removing most of the appliance-level switching "
    "noise. Fifth, a centred 7-day rolling mean is applied to suppress residual day-of-week "
    "variability while keeping the weekly and seasonal cycles intact."
)
add_para(
    "Missing days, where the rolling mean would otherwise produce a NaN, are forward and backward "
    "filled before the dataset is split temporally. The split keeps the first 80 percent of days "
    "for training and the last 20 percent for testing, which is the standard protocol for time "
    "series with non-stationary trends."
)

add_heading("3.2 ARIMA Model", level=2)
add_para(
    "ARIMA(p, d, q) is the workhorse of statistical time series forecasting. We follow the Box-"
    "Jenkins procedure. The Augmented Dickey-Fuller test on the training set returned a p-value "
    "of 0.94, indicating non-stationarity, so a single differencing order (d = 1) was applied. "
    "A grid search over p in [0, 2] and q in [0, 2] then selected the configuration with the "
    "lowest AIC, which was ARIMA(2, 1, 2) at AIC = 4547.63."
)
add_para(
    "The final ARIMA model is deployed under a rolling one-step-ahead protocol: at each test "
    "day, the model is re-fitted on all data observed up to that point and asked to forecast the "
    "next day only. This matches how a real operator would use such a model in production, where "
    "yesterday's actual reading is always available before today's forecast must be made.",
    lead_bold="Rolling forecast protocol. "
)

add_heading("3.3 LSTM Model", level=2)
add_para(
    "The LSTM model is a three-layer stacked recurrent network with 128, 64, and 32 units, "
    "interleaved with dropout layers at rate 0.2 and topped with a 16-unit ReLU dense layer "
    "before the linear output. Inputs are 14-day sliding windows of Min-Max scaled power values. "
    "Training uses the Adam optimiser with mean squared error loss, batch size 16, and up to 100 "
    "epochs with early stopping triggered after 5 stagnant validation epochs. A fixed random "
    "seed (42) is used to make the run reproducible across machines."
)

add_figure("lstm_architecture.png",
           "Figure 3. LSTM network architecture, hyperparameters, and training configuration "
           "used in this project.",
           width=Inches(5.5))

add_heading("3.4 FB-Prophet Model", level=2)
add_para(
    "FB-Prophet decomposes a series into trend, seasonality, and holiday effects, fitting each "
    "component as an additive (or multiplicative) signal. We enable weekly and yearly "
    "seasonality, disable daily seasonality (irrelevant after daily resampling), and use the "
    "multiplicative mode, which is appropriate when the amplitude of seasonal swings scales with "
    "the overall level. The changepoint prior scale is left at 0.05, the package default for "
    "moderately flexible trends."
)
add_para(
    "Crucially, Prophet is run in single multi-step mode, that is, it forecasts the entire "
    "122-day test horizon in one shot, without access to test-set actuals. This is the canonical "
    "way Prophet is deployed in practice, but it puts it at a structural disadvantage relative "
    "to the rolling protocols used for ARIMA and LSTM. We return to this point in the discussion.",
    lead_bold="Multi-step protocol. "
)

add_figure("rolling_forecast.png",
           "Figure 4. Comparison of the two forecasting protocols. ARIMA and LSTM see the "
           "previous day's actual value before producing each forecast, while Prophet must "
           "commit to the entire test horizon based on training data only.",
           width=Inches(6.3))

add_heading("3.5 Persistence Baseline", level=2)
add_para(
    "Before claiming that a complex model is useful, we must show that it beats the simplest "
    "possible alternative. We therefore include a persistence baseline that predicts the next day "
    "to equal the previous day's actual value (y_hat[t] = y[t-1]). This is the canonical sanity "
    "check for time series forecasting. Any model that cannot beat persistence has learned "
    "nothing beyond what a one-line rule already captures. Persistence runs under the same "
    "rolling protocol as ARIMA and LSTM: the previous actual day is always available when "
    "predicting the next."
)

add_heading("3.6 Evaluation Metrics", level=2)
add_para(
    "Predictive accuracy is measured with four complementary metrics. Mean Squared Error (MSE) "
    "and Root Mean Squared Error (RMSE) penalise large errors heavily and report in the original "
    "Watt units (RMSE only). Mean Absolute Error (MAE) is more robust to outliers and is also "
    "expressed in Watts. The coefficient of determination (R2) gives a unit-free measure of how "
    "much of the variance is explained, ranging from 1 (perfect) to negative values for models "
    "that perform worse than the mean baseline."
)

add_heading("3.7 Statistical Significance via Diebold-Mariano Test", level=2)
add_para(
    "Point estimates of RMSE or R2 can hide whether one model is genuinely better than another. "
    "We use the Diebold-Mariano (DM) test of equal predictive accuracy with the Harvey-Leybourne-"
    "Newbold small-sample correction. The null hypothesis is that two competing models have the "
    "same expected squared-error loss; a negative DM statistic favours the first model, positive "
    "the second. We report two-sided p-values for all pairwise comparisons of interest."
)

doc.add_page_break()

# ============================================================
# 4. EXPLORATORY DATA ANALYSIS
# ============================================================
add_heading("4. Exploratory Data Analysis", level=1)

add_para(
    "Before modelling, we inspected the cleaned aggregate signal to characterise its temporal "
    "structure. The four panels in Figure 5 summarise the analysis."
)

add_figure("data_analysis.png",
           "Figure 5. Exploratory analysis of House 2. Top-left: daily mean power across "
           "22 months. Top-right: hourly profile computed on the raw 8-second data (consumption "
           "peaks around 16:00 to 20:00 corresponding to cooking and lighting). Bottom-left: "
           "average daily power by weekday (weekends are higher). Bottom-right: distribution of "
           "daily power.",
           width=Inches(6.2))

add_para(
    "Three findings emerged. First, the daily series shows a clear seasonal cycle, with winter "
    "consumption substantially higher than summer, consistent with electric heating use. Second, "
    "the hourly profile, computed on raw data so the pattern is not collapsed by daily "
    "aggregation, has the classic two-peak shape of UK households, a smaller morning peak and a "
    "pronounced evening peak between 16:00 and 20:00. Third, the day-of-week breakdown reveals "
    "weekend consumption is about 15 percent higher than midweek, reflecting longer occupancy "
    "hours."
)

# ============================================================
# 5. RESULTS
# ============================================================
add_heading("5. Results", level=1)

add_heading("5.1 Final Performance Comparison", level=2)
add_para(
    "Table 1 reports the test-set metrics for the persistence baseline and the three models on "
    "House 2. The headline result is that ARIMA(2, 1, 2) is the winner with R2 = 0.876, but its "
    "edge over the trivial persistence baseline (R2 = 0.873) is razor-thin. LSTM (R2 = 0.748) and "
    "FB-Prophet (R2 = 0.015) both fall meaningfully below persistence. The implication, which we "
    "examine in detail in Section 5.5 with the Diebold-Mariano test, is that the lag-1 "
    "autocorrelation in daily residential energy is doing almost all of the predictive work; "
    "ARIMA captures it more cleanly, while LSTM and Prophet fail to recover it."
)

add_table_caption("Table 1. Test set performance on House 2 daily forecasts, including the "
                  "persistence baseline (y_hat[t] = y[t-1]).")

table = doc.add_table(rows=5, cols=5, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['Model', 'MSE', 'RMSE (W)', 'MAE (W)', 'R2']
data = [
    ['Persistence',    '2,426.76', '49.26', '33.51', '0.873'],
    ['ARIMA(2,1,2)',   '2,370.01', '48.68', '32.75', '0.876'],
    ['LSTM (3-layer)', '4,821.35', '69.44', '52.25', '0.748'],
    ['FB-Prophet',     '18,851.52', '137.30', '111.70', '0.015'],
]
for i, h in enumerate(headers):
    c = table.rows[0].cells[i]
    c.text = h
    for p in c.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(10)
            r.font.name = 'Calibri'
for r_idx, row in enumerate(data):
    for c_idx, v in enumerate(row):
        c = table.rows[r_idx + 1].cells[c_idx]
        c.text = v
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.size = Pt(10)
                r.font.name = 'Calibri'

doc.add_paragraph()

add_figure("metrics_comparison.png",
           "Figure 6. Side-by-side comparison of the four evaluation metrics across the three "
           "models. Lower is better for MSE, RMSE, and MAE; higher is better for R2.",
           width=Inches(6.3))

add_heading("5.2 Predicted versus Actual Series", level=2)
add_para(
    "Figure 7 plots each model's predictions against the held-out test values. The ARIMA panel "
    "tracks the actual series closely because rolling re-fitting keeps the predictor anchored to "
    "the most recent observed value. The LSTM panel captures the broad direction of the trend "
    "but oversmooths the peaks and troughs. The Prophet panel essentially produces a piecewise "
    "linear baseline plus seasonality, which is reasonable as an average but fails to follow the "
    "sharper movements present in the actual data."
)

add_figure("model_comparison.png",
           "Figure 7. Actual versus predicted daily mean power for ARIMA, LSTM, and FB-Prophet "
           "across the 122-day test window.",
           width=Inches(6.3))

add_heading("5.3 Training Diagnostics", level=2)
add_para(
    "The LSTM training loss curve in Figure 8 shows fast initial convergence and a plateau "
    "after roughly 12 epochs, at which point early stopping was triggered. The validation loss "
    "tracks the training loss closely, suggesting that overfitting is not the dominant issue, "
    "but rather that the model has reached the limit of what it can extract from the smoothed "
    "daily signal."
)

add_figure("lstm_training_history.png",
           "Figure 8. LSTM training and validation MSE loss across epochs.",
           width=Inches(5.0))

add_para(
    "Prophet's decomposition (Figure 9) highlights a clear yearly trend (consumption highest in "
    "winter, lowest in summer) and a moderate weekly cycle. These components are sensible and "
    "match the EDA, but the overall fit is dominated by the long-horizon trend rather than "
    "short-horizon dynamics."
)

add_figure("prophet_components.png",
           "Figure 9. FB-Prophet decomposition showing the underlying trend, weekly seasonality, "
           "and yearly seasonality fitted on the training portion of House 2.",
           width=Inches(5.5))

add_heading("5.4 Development Iterations", level=2)
add_para(
    "The final results in Table 1 were not achieved on the first attempt. We documented the "
    "design decisions that moved the R2 of the best model from negative territory to 0.876. "
    "Figure 10 summarises the three iterations."
)

add_figure("iteration_timeline.png",
           "Figure 10. Three iterations of the design and their effect on test R2. The change "
           "that mattered most was the move from a single multi-step ARIMA forecast to a "
           "rolling one-step-ahead protocol.",
           width=Inches(6.3))

doc.add_page_break()

# ============================================================
# 5.5 STATISTICAL SIGNIFICANCE
# ============================================================
add_heading("5.5 Statistical Significance: Diebold-Mariano Test", level=2)
add_para(
    "The pairwise Diebold-Mariano test on House 2 returns the values in Table 2. Three findings "
    "are worth flagging. First, ARIMA's apparent win over Persistence is not statistically "
    "significant at any conventional level (p = 0.665). The two models produce essentially "
    "indistinguishable forecasts on this test set. Second, both LSTM and Prophet are "
    "significantly worse than Persistence (p < 0.001 in both cases): adding model complexity "
    "actively harms accuracy here. Third, ARIMA does significantly beat LSTM and Prophet "
    "(p < 0.001 for both), so it is a reasonable choice when the alternatives are LSTM or Prophet, "
    "but a one-line lag-1 rule is just as good."
)

add_table_caption("Table 2. Diebold-Mariano test results on House 2 with the Harvey-Leybourne-"
                  "Newbold small-sample correction. n = 123 test days, two-sided p-values.")

table2 = doc.add_table(rows=6, cols=4, style='Light Grid Accent 1')
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
dm_headers = ['Comparison (Model 1 vs Model 2)', 'DM Statistic', 'p-value', 'Conclusion']
dm_data = [
    ['ARIMA vs Persistence',  '-0.433',  '0.665',     'No significant difference'],
    ['LSTM vs Persistence',   '+3.936',  '< 0.001',   'Persistence significantly better'],
    ['Prophet vs Persistence','+7.949',  '< 0.001',   'Persistence significantly better'],
    ['ARIMA vs LSTM',         '-3.765',  '< 0.001',   'ARIMA significantly better'],
    ['ARIMA vs Prophet',      '-7.957',  '< 0.001',   'ARIMA significantly better'],
]
for i, h in enumerate(dm_headers):
    c = table2.rows[0].cells[i]
    c.text = h
    for p in c.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(10)
            r.font.name = 'Calibri'
for r_idx, row in enumerate(dm_data):
    for c_idx, v in enumerate(row):
        c = table2.rows[r_idx + 1].cells[c_idx]
        c.text = v
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.size = Pt(10)
                r.font.name = 'Calibri'

doc.add_paragraph()

add_figure("dm_test_plot.png",
           "Figure 11. Diebold-Mariano statistics for House 2. Green bars are significant at the "
           "5 percent level (|DM| > 1.96); grey bars are not. The ARIMA versus Persistence "
           "comparison is the only non-significant pair.",
           width=Inches(6.0))

# ============================================================
# 5.6 CROSS-HOUSE VALIDATION
# ============================================================
add_heading("5.6 Cross-House Validation", level=2)
add_para(
    "To check that the House 2 ranking is not an artefact of one household, we repeated the "
    "entire pipeline on House 1 and House 5. These two homes have different occupancies and "
    "appliance mixes (House 1 has freezers and a tumble dryer dominating the load, House 5 has a "
    "much higher mean daily power), so any ranking that survives across all three is unlikely to "
    "be a fluke. Table 3 reports the per-house R2 values."
)

add_table_caption("Table 3. R2 of each model across three REFIT households. Persistence is "
                  "the trivial baseline; ARIMA wins narrowly on every house, while LSTM and "
                  "Prophet underperform persistence consistently.")

table3 = doc.add_table(rows=5, cols=5, style='Light Grid Accent 1')
table3.alignment = WD_TABLE_ALIGNMENT.CENTER
ch_headers = ['Model', 'House 1', 'House 2', 'House 5', 'Verdict']
ch_data = [
    ['Persistence', '0.963', '0.873', '0.901', 'Strong baseline'],
    ['ARIMA',       '0.976', '0.876', '0.905', 'Best on every house'],
    ['LSTM',        '0.859', '0.748', '0.647', 'Below persistence on every house'],
    ['FB-Prophet',  '-2.315', '0.015', '-1.186', 'Often negative R2'],
]
for i, h in enumerate(ch_headers):
    c = table3.rows[0].cells[i]
    c.text = h
    for p in c.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(10)
            r.font.name = 'Calibri'
for r_idx, row in enumerate(ch_data):
    for c_idx, v in enumerate(row):
        c = table3.rows[r_idx + 1].cells[c_idx]
        c.text = v
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.size = Pt(10)
                r.font.name = 'Calibri'

doc.add_paragraph()

add_figure("cross_house_comparison.png",
           "Figure 12. Cross-house R2 by model. The ARIMA versus Persistence gap stays small on "
           "all three houses, LSTM is consistently below Persistence, and Prophet produces "
           "negative R2 on House 1 and House 5.",
           width=Inches(6.3))

add_para(
    "Two findings carry forward. First, the ARIMA-Persistence gap is small on every house, so the "
    "House 2 result is not unusual. Second, Prophet's R2 is negative on two of the three houses, "
    "indicating that it predicts worse than the unconditional mean. This is a stronger statement "
    "than the House 2 result alone suggested."
)

# ============================================================
# 5.7 PRACTICAL INTERPRETATION
# ============================================================
add_heading("5.7 Practical Interpretation of the Errors", level=2)
add_para(
    "RMSE values are easier to understand once expressed in domain units. For House 2, the mean "
    "test consumption is 434.8 W, so ARIMA's RMSE of 48.7 W is roughly 11.2 percent of the typical "
    "daily level. Translated to energy, an average-power error of 48.7 W sustained over 24 hours "
    "equals about 1.17 kWh per day. At the UK 2025 domestic electricity unit price of roughly "
    "0.27 GBP per kWh, this is an expected billing error of about 0.32 GBP per day, or about "
    "10 GBP per month for the single household. In an aggregator setting that prices day-ahead "
    "imbalance at, say, 0.10 GBP per kWh, the cost falls to about 0.12 GBP per day. These are "
    "small absolute numbers per household, but they scale linearly with portfolio size."
)

doc.add_page_break()

# ============================================================
# 6. DISCUSSION
# ============================================================
add_heading("6. Discussion", level=1)

add_heading("6.1 The Persistence Result, and What It Tells Us", level=2)
add_para(
    "The most important finding of this study is not that ARIMA beats the other models, but that "
    "ARIMA barely beats persistence and the gap is not statistically significant. Persistence "
    "predicts that tomorrow equals today: a one-line rule with zero learned parameters. On "
    "House 2 it reaches R2 = 0.873, only 0.003 below ARIMA's 0.876. The Diebold-Mariano test "
    "(p = 0.665) confirms that the two models have indistinguishable squared-error losses. "
    "The cross-house results paint the same picture: ARIMA's R2 edge over Persistence is 0.013 "
    "on House 1 and 0.004 on House 5."
)
add_para(
    "The mechanical reason is the strong lag-1 autocorrelation of 7-day-smoothed daily "
    "consumption. Yesterday is an excellent predictor of today, and any model that uses yesterday "
    "directly will perform almost identically. ARIMA(2, 1, 2) adds two more lags and two moving-"
    "average terms, but the extra information is small compared to lag-1. The practical implication "
    "is clear: at this granularity and on this signal, do not deploy heavy machinery when a "
    "one-line rule already captures the predictable structure."
)

add_heading("6.2 Why LSTM Underperforms the Trivial Baseline", level=2)
add_para(
    "LSTM not only loses to ARIMA, it loses significantly to Persistence (p < 0.001 on House 2, "
    "and a similar pattern on the other two houses). This is the more damning result: a 128,929-"
    "parameter neural network is worse than predicting that tomorrow equals today. The training "
    "set contains 490 daily observations, which gives only 476 input sequences after the 14-day "
    "lookback. With this little data, the network has too much capacity, learns smooth average "
    "behaviour, and oversmooths the day-to-day movements that Persistence captures exactly by "
    "construction. Recent work (Gasparin and others, 2024) reports a similar pattern: below six "
    "months of training data, simple baselines beat deep models. Our data is longer but the "
    "smoothed daily signal apparently still falls in the regime where deep learning is a net "
    "negative."
)

add_heading("6.3 Why Prophet Looks the Worst", level=2)
add_para(
    "Prophet's near-zero R2 on House 2 was not encouraging. Cross-house validation makes the "
    "picture worse: R2 = -2.315 on House 1 and -1.186 on House 5, meaning Prophet predicts "
    "considerably worse than the unconditional training mean. Two factors are at play. First, "
    "Prophet produces a single multi-step forecast for the entire test horizon with no access "
    "to test actuals, which puts it at a structural disadvantage against the rolling Persistence, "
    "ARIMA, and LSTM setups. Second, Prophet's strength lies in capturing slow seasonal trends "
    "and calendar effects; on smoothed daily residential data the dominant signal is the "
    "previous day, which Prophet's decomposition does not target. A fair comparison would also "
    "run Prophet in a rolling refit mode, but the structural mismatch with the data character "
    "would remain."
)

add_heading("6.4 Limitations and Threats to Validity", level=2)
add_para(
    "Three limitations qualify our conclusions."
)
add_bullet("Three households, not twenty. Cross-house validation strengthens the original "
           "single-house finding but does not yet test on all 17 remaining REFIT homes.")
add_bullet("Univariate setup. None of the models receive weather, calendar, or appliance-level "
           "exogenous data. Adding these could shift the picture, particularly for LSTM, which is "
           "designed to absorb richer inputs.")
add_bullet("Heavy smoothing. The 7-day rolling mean removes much of the high-frequency content. "
           "An operational forecaster who needs daily ramping information would have to revisit "
           "this choice, and the persistence advantage may shrink at finer time scales.")

add_heading("6.5 Comparison with Reference Literature", level=2)
add_para(
    "The ordering ARIMA greater than LSTM greater than Prophet, observed here, matches the "
    "findings of Bülüç, Sevli, and Yünlü (2025), who reported R2 = 0.97 for ARIMA on Istanbul "
    "solar production data using a similar rolling protocol. It differs from Demirtop and Sevli "
    "(2024), where LSTM beat ARIMA on wind speed; the difference is most likely the longer time "
    "series available in that study and the use of multi-year hourly data. Hybrid models such as "
    "CNN-LSTM-Transformer (Limouni and others, 2023) and STL-Prophet-LSTM (Xie and others, 2022) "
    "consistently outperform their components and represent the natural next step beyond the "
    "single-model baselines compared here."
)

# ============================================================
# 7. CONCLUSION
# ============================================================
add_heading("7. Conclusion and Future Work", level=1)

add_para(
    "We built a reproducible pipeline that forecasts the daily mean power consumption of three "
    "UK households using four methods, including a persistence baseline. Under a rolling one-"
    "step-ahead protocol and a fixed random seed, ARIMA(2, 1, 2) was the best model on every "
    "house, but its margin over a one-line persistence rule was not statistically significant "
    "(Diebold-Mariano p = 0.665 on House 2 and similarly small differences on Houses 1 and 5). "
    "LSTM and FB-Prophet both lost significantly to persistence in all three settings, with "
    "Prophet producing negative R2 on two of the three houses. The lesson we carry forward is "
    "that residential daily energy at this granularity is dominated by lag-1 autocorrelation, "
    "and that adding model capacity over a trivial baseline produces no measurable gain (ARIMA) "
    "or actively hurts (LSTM, Prophet)."
)

add_para(
    "The methodological lesson is just as important: reporting only complex models against each "
    "other can hide the fact that the entire model class is unnecessary. A persistence baseline "
    "and a Diebold-Mariano significance test together turn what looked like a clean ARIMA win "
    "into a more honest 'the data is too predictable for the models to matter' result."
)

add_para(
    "Four directions stand out for further work."
)
add_bullet("Repeating the analysis on a higher-frequency target (hourly or 15-minute) where "
           "lag-1 autocorrelation is weaker and the deep models may genuinely add value.")
add_bullet("Adding weather (temperature, irradiance) and calendar features through SARIMAX and "
           "multivariate LSTM, which give the heavier models information persistence cannot use.")
add_bullet("Extending the cross-house comparison from three to all 20 REFIT homes, and exploring "
           "cross-household transfer learning.")
add_bullet("Hybrid models such as STL-Prophet-LSTM and CNN-LSTM-Transformer that combine the "
           "strengths of decomposition and neural sequence learning.")

add_heading("Reproducibility Note", level=2)
add_para(
    "All experiments use a fixed random seed (42) for NumPy, Python's random module, and "
    "TensorFlow. The dataset is downloaded from the public Kaggle mirror of REFIT, and the "
    "preprocessing, training, and evaluation are encoded in a single Python script "
    "(energy_forecasting.py). Results in this report were obtained on Python 3.13, TensorFlow "
    "2.x, statsmodels 0.14, and Prophet 1.3."
)

# ============================================================
# REFERENCES
# ============================================================
add_heading("References", level=1)

refs = [
    "Atalay, B. A., and Zor, K. (2025). XGBoost ile hidroelektrik enerji tahmini. "
    "Çukurova Üniversitesi Mühendislik Fakültesi Dergisi, 40(1), 205-218.",

    "Berus, Y., and Yakut, Y. B. (2024). Derin öğrenme (1D-CNN, RNN, LSTM, BiLSTM) ile enerji "
    "tüketim tahmini: Diyarbakır AVM örneği. Dicle University Journal of Engineering, 15(2), "
    "311-322.",

    "Bülüç, M., Sevli, O., and Yünlü, L. (2025). Time Series Analysis of Solar Energy Production "
    "Based on Weather Conditions. GU J Sci, Part A, 12(4), 1060-1077.",

    "Çolak, M. B., and Özhan, E. (2025). Renewable energy forecasting in Turkey: Analytical "
    "approaches. Journal of Intelligent Systems: Theory and Applications, 8(1), 25-34.",

    "Demirtop, A., and Sevli, O. (2024). Wind speed prediction using LSTM and ARIMA time series "
    "analysis models: A case study of Gelibolu. Turkish Journal of Engineering, 8(3), 524-536.",

    "Gasparin, A., Lukovic, S., and Alippi, C. (2024). Load Forecasting for Households and "
    "Energy Communities: Are Deep Learning Models Worth the Effort? arXiv:2501.05000.",

    "Limouni, T., Yaagoubi, R., Bouziane, K., Guissi, K., and Baali, E. H. (2023). Solar Energy "
    "Production Forecasting Based on a Hybrid CNN-LSTM-Transformer Model. Mathematics, 11(3), 676.",

    "Murray, D., Stankovic, L., and Stankovic, V. (2017). An electrical load measurements "
    "dataset of United Kingdom households from a two-year longitudinal study. Scientific Data, "
    "4, 160122.",

    "Rahman, M. R., and others (2025). Time-series and deep learning approaches for renewable "
    "energy forecasting in Dhaka: a comparative study of ARIMA, SARIMA, and LSTM models. "
    "Discover Sustainability, 6, 01733.",

    "Sakib, N., and others (2024). Deep learning-driven hybrid model for short-term load "
    "forecasting and smart grid information management. Scientific Reports, 14, 63262.",

    "Tamay, M., and Türker, G. F. (2024). Machine learning based energy forecasting for "
    "photovoltaic solar plants. Yekarum, 9(2), 128-146.",

    "Triebe, O., Hewamalage, H., Pilyugina, P., Laptev, N., Bergmeir, C., and Rajagopal, R. "
    "(2021). NeuralProphet: Explainable Forecasting at Scale. arXiv:2111.15397.",

    "Vargas-Forero, V. M., Manotas-Duque, D. F., and Trujillo, L. (2025). Comparative study of "
    "forecasting methods to predict the energy demand for the market of Colombia. International "
    "Journal of Energy Economics and Policy, 15(1), 65-76.",

    "Xie, J., Li, Z., Zhou, Z., and Liu, S. (2022). A hybrid forecasting model using LSTM and "
    "Prophet for energy consumption with decomposition of time series data. PeerJ Computer "
    "Science, 8, e1001.",
]

for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.first_line_indent = Cm(-0.6)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f"[{i}] {ref}")
    r.font.size = Pt(9.5)
    r.font.name = 'Calibri'


# ============================================================
# SAVE DOCX
# ============================================================
doc.save(OUTPUT_DOCX)
print(f"DOCX saved: {OUTPUT_DOCX}")
print(f"DOCX size: {os.path.getsize(OUTPUT_DOCX)/1024:.0f} KB")

# ============================================================
# CONVERT TO PDF
# ============================================================
print("\nConverting to PDF...")
try:
    from docx2pdf import convert
    convert(OUTPUT_DOCX, OUTPUT_PDF)
    print(f"PDF saved: {OUTPUT_PDF}")
    print(f"PDF size: {os.path.getsize(OUTPUT_PDF)/1024:.0f} KB")
except Exception as e:
    print(f"docx2pdf failed: {e}")
    print("Will try alternative PDF conversion methods.")
