"""
Generate the comprehensive Final Project Report (DOCX) and convert to PDF.

Clean style:
 - Calibri throughout
 - 1.3 line spacing
 - No em-dash, no superscript R-squared, no et al., no subscript digits
 - "and others" instead of "et al."

Rich content:
 - Cover page
 - Abstract
 - Section structure: Intro -> Data -> Methodology -> Results -> Discussion -> Conclusion
 - All 6 EDA/result figures + 4 custom diagrams
 - Results tables
 - References
"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.abspath(__file__))
RESULTS = os.path.join(BASE, "results")
OUTPUT_DOCX = os.path.join(BASE, "EBT629E_Final_Project_Report_TR.docx")
OUTPUT_PDF = os.path.join(BASE, "EBT629E_Final_Project_Report_TR.pdf")

doc = Document()

# Page setup, tighter margins and document-level auto-hyphenation
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

# Enable Word auto-hyphenation (fills justified lines without huge gaps)
settings_xml = doc.settings.element
auto_h = OxmlElement('w:autoHyphenation')
auto_h.set(qn('w:val'), 'true')
settings_xml.append(auto_h)

# Default style
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.line_spacing = 1.3
style.paragraph_format.space_after = Pt(5)

NAVY = RGBColor(0, 51, 102)
GREY = RGBColor(80, 80, 80)


# ============================================================
# HELPER FUNCTIONS
# ============================================================
def add_heading(text, level=1, color=NAVY):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
        run.font.name = 'Calibri'
    return h


def add_para(text, lead_bold=None, justify=True):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if lead_bold:
        r = p.add_run(lead_bold)
        r.bold = True
        r.font.name = 'Calibri'
        r.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.name = 'Calibri'
    r2.font.size = Pt(11)
    return p


def add_figure(filename, caption, width=Inches(6.0)):
    path = os.path.join(RESULTS, filename)
    if not os.path.exists(path):
        doc.add_paragraph(f"[Figure missing: {filename}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run()
    run.add_picture(path, width=width)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    r = cap.add_run(caption)
    r.bold = True
    r.font.size = Pt(9)
    r.font.name = 'Calibri'


def add_bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    for r in p.runs:
        r.font.name = 'Calibri'
        r.font.size = Pt(11)


def add_table_caption(text):
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(2)
    r = cap.add_run(text)
    r.bold = True
    r.font.size = Pt(9)
    r.font.name = 'Calibri'


# ============================================================
# COVER PAGE
# ============================================================
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("İstanbul Teknik Üniversitesi")
r.font.size = Pt(14)
r.font.name = 'Calibri'
r.font.color.rgb = GREY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("EBT 629E, Yapay Zeka")
r.font.size = Pt(14)
r.font.name = 'Calibri'
r.font.color.rgb = GREY

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Dönem Sonu Proje Raporu")
r.font.size = Pt(22)
r.bold = True
r.font.name = 'Calibri'
r.font.color.rgb = NAVY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Zaman Serisi Modelleri ile\nHane Halkı Enerji Tüketimi Tahmini")
r.font.size = Pt(16)
r.bold = True
r.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("REFIT Veri Seti Üzerinde ARIMA, LSTM ve FB-Prophet Modellerinin Karşılaştırmalı Analizi")
r.font.size = Pt(12)
r.italic = True
r.font.name = 'Calibri'
r.font.color.rgb = GREY

for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Grup Üyeleri")
r.font.size = Pt(12)
r.bold = True
r.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    "Nurullah Yıldırım (301252004)\n"
    "Kadir Göksel Gündüz (301241077)\n"
    "Furkan Çınar (301212001)"
)
r.font.size = Pt(12)
r.font.name = 'Calibri'

for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Mayıs 2026")
r.font.size = Pt(12)
r.font.name = 'Calibri'
r.font.color.rgb = GREY

doc.add_page_break()

# ============================================================
# ABSTRACT
# ============================================================
add_heading("Özet", level=1)
add_para(
    "Bu rapor, REFIT Akıllı Ev veri setinden hane halkı enerji tüketimine uygulanan dört zaman "
    "serisi tahmin yaklaşımının karşılaştırmalı bir incelemesini sunmaktadır. Üç Birleşik Krallık "
    "hanesinden (Ev 1, Ev 2 ve Ev 5) 22 aylık 8 saniye çözünürlüklü ölçümleri kullanıyoruz; "
    "verileri günlük ortalama gücü üretecek şekilde yeniden örneklenir, ardından 7 günlük hareketli "
    "ortalama ile yumuşatılır ve her seri yüzde 80 eğitim, yüzde 20 test olacak şekilde zaman "
    "sırasında bölünür. Üç aday modelin (ARIMA, Long Short-Term Memory (LSTM) ve FB-Prophet) yanı "
    "sıra sağlık kontrolü olarak persistence temel modeli (y_hat[t] = y[t-1]) ekledik. ARIMA modeli "
    "ADF durağanlık testi ve AIC ızgara araması ile ayarlanır, ardından tek adım önde yuvarlanan "
    "tahmin protokolü altında çalıştırılır; LSTM üç katmanlı yığılmış RNN mimarisi kullanır; "
    "Prophet haftalık ve yıllık mevsimsellik bileşenleriyle tek seferlik çok adım modunda çalışır. "
    "Ev 2 de ARIMA(2, 1, 2) R2 = 0.876 ve RMSE = 48.7 W e ulaşır, ancak Diebold-Mariano testi "
    "persistence temel modeli (R2 = 0.873) üzerindeki marjinin istatistiksel olarak anlamlı "
    "olmadığını gösterir (p = 0.665). LSTM (R2 = 0.748) ve Prophet (R2 = 0.015) persistence e "
    "anlamlı olarak kaybeder (her ikisinde de p < 0.001). Aynı örüntü Ev 1 ve Ev 5 te de tekrarlanir: "
    "ARIMA dar bir farkla kazanır, ancak persistence üzerine çok az sey ekler; LSTM ve Prophet ise "
    "tutarlı olarak temel modelin altında kalır. Baslica ders, yumuşatılmış günlük konut tüketiminin "
    "gecikme-1 otokorelasyonu tarafından domine edildiği; bu nedenle tek satırlık bir kuralın "
    "neredeyse tüm tahmin edilebilir yapıyı yakaladığıdır. Daha ağır makineler ancak daha güçlü "
    "dışsal girdiler veya daha ince granulariteli hedefler bu otokorelasyonu kırdığında doğrulanır. "
    "Sabit rastgele tohum kullanılan tüm kod tabanı bu rapora eşlik etmektedir."
)

doc.add_page_break()

# ============================================================
# 1. INTRODUCTION
# ============================================================
add_heading("1. Giriş", level=1)

add_heading("1.1 Genel Bilgi ve Motivasyon", level=2)
add_para(
    "Hane halkı elektrik tüketimi, modern enerji sistemlerinin merkezindeki birkaç sorunu aynı anda etkilemektedir. Şebeke işletmecisi açısından kısa ufuklu doğrulukla yapılan tahminler, yenilenebilir uretiminin dengelenmesi, yedek kapasite planlamasi ve kısıtlama miktarinin azaltılması için gereklidir. Tuketici açısından ise bu tahminler talep yaniti programlarına katilimi, cihazlarin akıllı şekilde programlanmasını, çatı üstü PV ve batarya sistemleriyle entegrasyonun iyileştirilmesini kolaylastirir. Avrupa genelinde akıllı sayaçların hizla yayilmasi ve makine ogrenmesi yontemlerinin olgunlaşması, sorunun artık trafo merkezi yerine tekil hane düzeyinde çalışılmasını mumkun kilmaktadir."
)
add_para(
    "Konut tüketiminin tahmini, toplulastirilmis talebi tahmin etmekten yapisal olarak daha zordur. Tek bir hane ani anahtarlama olaylari (su isiticisi, firin, çamaşır makinesi), kullanıcıya özgü programlar ve hava ile tatil etkilesimi gösteren mevsimsellik içerir. Çok sayida haneyi ortalayarak bu gürültünün büyük bolumu sonuç sinyalinden silinir; bu nedenle şebeke düzeyinde düşük MAPE raporlanirken hane düzeyinde tahmin tamamen başarısız olabilir."
)

add_heading("1.2 Proje Hedefleri", level=2)
add_para(
    "Bu projenin sordugu soru hem odakli hem de uygulamaya yoneliktir: yaygin üç zaman serisi modeli olan ARIMA, LSTM ve FB-Prophet arasında, tek bir hanenin günlük enerji tuketimini tahmin etmek için hangisi daha uygundur ve bunun ardındaki gerekçeler nelerdir? Somut olarak proje dört hedef pesinden gider."
)
add_bullet("Ham 8 saniye REFIT okumalarindan model ciktilarina kadar uzanan tekrar üretilebilir bir boru hatti kurmak.")
add_bullet("ARIMA, LSTM ve FB-Prophet modellerini aynı veri bölme protokolü altında tek bir test ortaminda karsilastirmak.")
add_bullet("Model performansini MSE, RMSE, MAE ve R2 metrikleriyle nicel olarak olcmek.")
add_bullet("Her modelin nerede kazanip nerede kaybettigini incelemek ve nihai tasarima goturen iterasyonları belgelemek.")

add_heading("1.3 Rapor Yapisi", level=2)
add_para(
    "Bölüm 2, REFIT veri setini ve seçilen haneyi tanitir. Bölüm 3, ön işleme boru hatti ve üç modelleme yaklasimini detaylandirir; literaturdeki çok çeşitli karsilastirmalardan ayrışmamızı sağlayan yuvarlanan tahmin protokolü de burada anlatılır. Bölüm 4 keşifsel veri analizini sunar. Bölüm 5 sayisal sonuçları ve görsel karsilastirmalari sunar. Bölüm 6 bulgulari, çalışmanın sinirlamalarini ve geçerlilik tehditlerini tartisir. Bölüm 7 sonuç ve ileride yapilabilecek çalışmalarla kapanır."
)

# ============================================================
# 2. DATASET
# ============================================================
add_heading("2. Veri Seti: REFIT Akıllı Ev Elektrik Yükü Ölçümleri", level=1)

add_heading("2.1 Genel Bakış", level=2)
add_para(
    "REFIT (Personalised Retrofit Decision Support Tools for UK Homes), Ekim 2013 ile Haziran 2015 arasında Birleşik Krallık Loughborough bolgesinde toplanan boylamsal bir veri setidir. Veriler Strathclyde, Loughborough ve East Anglia Üniversiteleri ortaklığında EPSRC finansmanıyla üretilmiş ve Creative Commons Attribution 4.0 International lisansı altında yayinlanmistir. Yirmi hane (Ev 1 ile Ev 21 arasında, Ev 14 atlanarak) tüm hane akımı olcen bir akım klempi ve dokuz adet bireysel cihaz olcumleyici ile donatilmistir. Aktif güç, yaklaşık altı ile sekiz saniyede bir Watt cinsinden örneklenmiş ve tüm haneler toplaminda yaklaşık 1,19 milyar gözlemden oluşan bir koleksiyon oluşturmuştur (Murray, Stankovic ve Stankovic, 2017)."
)

add_heading("2.2 Seçilen Hane: Ev 2", level=2)
add_para(
    "Bu projenin test ortami olarak Ev 2 seçilmiştir. CSV dosyasi diskte 299 MB yer kaplar ve 22 aylık bir surede toplanan 5.733.526 zaman damgali okuma içerir. Izlenen cihazlar Buzdolabi-Dondurucu, Çamaşır Makinesi, Bulasik Makinesi, Televizyon, Mikrodalga, Tost Makinesi, Hi-Fi, Su Isiticisi ve Firin Aspiratorudur. Ev 2; belgelenmiş cihaz degisikligi olmamasi, çatı üstü PV girişimi bulunmamasi ve tipik bir Birleşik Krallık hanesinin cihaz karisimini temsil etmesi nedeniyle seçilmiştir."
)

# Correlation heatmap
add_figure("correlation_heatmap.png",
           "Şekil 1. Ev 2 hanesinde toplam güç kanali ile dokuz bireysel cihaz olcumleyici "
           "arasındaki Pearson korelasyonu.",
           width=Inches(5.5))

add_heading("2.3 Veri Kalitesi ve Temizleme", level=2)
add_para(
    "REFIT surumu, kısa bosluklarda ileri doldurma uygular ve iki dakikadan uzun bosluklari sifirlar. Cihaz kanallarinda 4000 W üzerindeki okumalar (sensor menzilinin dışında) orijinal bakımcılar tarafından zaten kaldirilmistir. Bizim ek temizleme adımlarımız, negatif toplam okumalari (klemp polarite hatalarından kaynaklanan) kaldirir ve degerlerin ust yüzde 1 lik dilimini aykırı olarak budar. Bu adımlardan sonra toplam seri sureklidir ve yeniden örnekleme için hazirdir."
)

doc.add_page_break()

# ============================================================
# 3. METHODOLOGY
# ============================================================
add_heading("3. Yöntem", level=1)

add_figure("pipeline_diagram.png",
           "Şekil 2. Uctan uca makine ogrenmesi boru hatti. Ham 8 saniye REFIT verisi aykırı "
           "değer temizleme, günlük yeniden örnekleme ve 7 günlük yumuşatma adımlarından gecer, "
           "ardından üç rakip modeli besler. Modellerin tahminleri ayrılmış test seti üzerinde "
           "değerlendirilir.",
           width=Inches(6.3))

add_heading("3.1 Veri Ön İşleme", level=2)
add_para(
    "Bes ön işleme adımı ham sinyali modellemeye hazırlar. Önce, zaman damgasi sutunu pandas DatetimeIndex e çevrilir. Ikinci olarak, bellek tasarrufu için Unix epoch sutunu bırakılır. Ucuncu olarak, aykırı değer filtreleme negatif değerleri ve 99 uncu persentilin üzerindeki girdileri kaldirir. Dordunculuk olarak, temizlenmis seri günlük ortalamaya yeniden örneklenir; bu islem 5,7 milyon satiri yaklaşık 630 günlük gozleme indirir ve cihaz duzeyindeki anahtarlama gürültüsünün cogunu temizler. Besinci olarak, ortalanmis 7 günlük hareketli ortalama, hafta-ici değişkenliği bastırırken haftalık ve mevsimsel donguleri korur."
)
add_para(
    "Hareketli ortalamanin NaN ürettiği eksik günler, veri seti zaman sırasıyla bolunmeden önce ileri ve geri doldurma ile tamamlanir. Bolme ilk yüzde 80 i eğitim, son yüzde 20 yi test olarak ayırır; bu, mevsimsel kayma içeren zaman serileri için standart protokoldur."
)

add_heading("3.2 ARIMA Modeli", level=2)
add_para(
    "ARIMA(p, d, q), istatistiksel zaman serisi tahmininin temel modelidir. Box-Jenkins prosedurunu izledik. Augmented Dickey-Fuller testi eğitim setinde 0.94 p-değerine sahip olarak durağansızlığa isaret etti; bu nedenle birinci dereceden farklılaştırma (d = 1) uygulandi. Ardından p in [0, 2] ve q in [0, 2] üzerinde ızgara araması yapilarak en düşük AIC değerine sahip yapılanma secildi: ARIMA(2, 1, 2), AIC = 4547.63."
)
add_para(
    "Nihai ARIMA modeli yuvarlanan tek-adım-önde tahmin protokolünde çalıştırılır: her test gununde model, o ana kadar gozlenen tüm veriler üzerinde yeniden uydurulur ve sadece bir sonraki gün için tahmin yapar. Bu, bir operatorun modelden uretimde nasıl yararlanacagini birebir yansıtır; çünkü dünün gercek olcumu bugünün tahmininden önce her zaman elimizdedir.",
    lead_bold="Yuvarlanan tahmin protokolü. "
)

add_heading("3.3 LSTM Modeli", level=2)
add_para(
    "LSTM modeli, sırasıyla 128, 64 ve 32 hucreli üç katmanlı yığılmış tekrarlayan ağdır; katmanlar arasında 0.2 oraninda dropout, en uste ise 16 hucreli ReLU yogun katman ve ardından dogrusal çıkış bulunur. Girdiler, 14 günlük kayar pencerelere uygulanan Min-Max ölçekli güç degerleridir. Egitimde Adam iyilestirici, MSE kayip fonksiyonu, 16 yigin büyüklüğü ve 100 epoch kullanıldı; 5 epoch boyunca doğrulama kaybinda iyileşme olmazsa eğitim erken durdurulur. Sabit rastgele tohum (42), çalışmanın makineler arasi tekrar üretilebilir olmasini sağlar."
)

add_figure("lstm_architecture.png",
           "Şekil 3. Bu projede kullanılan LSTM ağının mimarisi, hiperparametreleri ve eğitim "
           "yapilandirmasi.",
           width=Inches(5.5))

add_heading("3.4 FB-Prophet Modeli", level=2)
add_para(
    "FB-Prophet, bir seriyi trend, mevsimsellik ve tatil etkilerine ayrıştırır ve her bileşeni ek (veya carpan) sinyal olarak uydurur. Haftalık ve yıllık mevsimselliği etkinleştirdik, günlük mevsimselliği devre disi bıraktık (günlük yeniden örnekleme sonrasinda anlamsızdır) ve carpan modu kullandik; çünkü mevsimsel salinim genligi genel düzeyle ölçeklenir. Degisim noktasi onsel olcegi varsayilan 0.05 te bırakıldı."
)
add_para(
    "Önemli olan, Prophet in tek seferlik çok-adım modunda çalıştırılmasıdır: yanı test ufkunun tamamini (122 gün) bir kerede tahmin eder ve test seti gercek degerlerine erisemez. Bu, Prophet in pratikteki kanonik kullanimidir, ancak ARIMA ve LSTM in yuvarlanan kurulumlarina kiyasla onu yapisal olarak dezavantajli kilar. Tartismada bu noktaya geri donecegiz.",
    lead_bold="Çok-adım protokolü. "
)

add_figure("rolling_forecast.png",
           "Şekil 4. İki tahmin protokolunun karşılaştırması. ARIMA ve LSTM her tahmin öncesinde "
           "bir önceki günün gercek değerini gorur; Prophet ise yalnizca eğitim verisine "
           "dayanarak test ufkunun tamamina taahhut eder.",
           width=Inches(6.3))

add_heading("3.5 Persistence Temel Modeli", level=2)
add_para(
    "Karmasik bir modelin gercekten faydali oldugunu iddia etmeden önce, en basit alternatifi "
    "yenebildigini göstermemiz gerekir. Bu nedenle, bir sonraki günün bir önceki günün gercek "
    "değerine eşit olduğu tahminini (y_hat[t] = y[t-1]) üretmesi için persistence (yapiskanlik) "
    "temel modeli ekledik. Bu, zaman serisi tahmininde kanonik bir sağlık kontrolü modelidir. "
    "Persistence'i yenemeyen herhangi bir model, tek satırlık bir kuralın yakaladigindan fazla "
    "bir sey ogrenmemis demektir. Persistence, ARIMA ve LSTM ile aynı yuvarlanan protokol "
    "altında çalışır: bir sonraki günü tahmin ederken bir önceki günün gercek degeri her zaman "
    "elimizdedir."
)

add_heading("3.6 Değerlendirme Metrikleri", level=2)
add_para(
    "Tahmin doğruluğu birbirini tamamlayan dört metrikle olculur. Ortalama Karesel Hata (MSE) ve Karekok Ortalama Karesel Hata (RMSE) büyük hatalari ağır biçimde cezalandirir; RMSE ozgun birim olan Watt cinsinden raporlanir. Ortalama Mutlak Hata (MAE) aykırı değerlere daha dayanıklıdır ve yine Watt cinsindendir. Belirleme katsayisi (R2), aciklanan varyansin birimsiz olcumudur; 1 mükemmel tahmini, ortalama temel modelinin altında kalan modeller için negatif değerleri ifade eder."
)

add_heading("3.7 Diebold-Mariano Testi ile İstatistiksel Anlamlılık", level=2)
add_para(
    "RMSE veya R2 nin noktasal değerleri, bir modelin diğerinden gercekten daha iyi olup olmadığını "
    "gizleyebilir. Bu nedenle Harvey-Leybourne-Newbold küçük orneklem düzeltmesiyle Diebold-Mariano "
    "(DM) eşit tahmin doğruluğu testini kullanıyoruz. Bos hipotez, iki rakip modelin beklenen "
    "karesel hata kaybinin eşit oldugudur; negatif DM istatistigi birinci modeli, pozitif olan "
    "ikinci modeli destekler. Ilgilenilen tüm ikili karşılaştırmalar için iki tarafli p-degerlerini "
    "raporluyoruz."
)

doc.add_page_break()

# ============================================================
# 4. EXPLORATORY DATA ANALYSIS
# ============================================================
add_heading("4. Keşifsel Veri Analizi", level=1)

add_para(
    "Modellemeden önce, temizlenmis toplam sinyali zaman yapisini karakterize etmek için inceledik. Şekil 5 teki dört panel bu analizi özetler."
)

add_figure("data_analysis.png",
           "Şekil 5. Ev 2 nin keşifsel analizi. Sol ust: 22 ay boyunca günlük ortalama güç. "
           "Sag ust: ham 8 saniye verisi üzerinden hesaplanan saatlik profil (16:00 ile 20:00 "
           "arasında yemek pisirme ve aydınlatmaya bağlı zirve göze carpiyor). Sol alt: hafta "
           "günlerine göre ortalama günlük güç (hafta sonlari daha yüksek). Sag alt: günlük "
           "gucun dagilimi.",
           width=Inches(6.2))

add_para(
    "Üç bulgu one çıktı. Birincisi, günlük seri belirgin bir mevsimsel dongu sergiliyor: kis tüketimi yaza göre belirgin şekilde yüksek; bu, elektrikli isitma kullanimiyla uyumlu. İkincisi, günlük toplulastirmanin orneklemi cokertmemesi için ham veriden hesaplanan saatlik profil, Birleşik Krallık hanelerine özgü klasik iki tepeli şekli gösteriyor: küçük bir sabah tepesi ve 16:00 ile 20:00 arasında belirgin bir akşam tepesi. Üçüncüsü, hafta-icine kiyasla hafta sonu tüketiminin yaklaşık yüzde 15 daha yüksek olmasi, gün icindeki dolu sure uzunluguyla orustugu konutu yansitiyor."
)

# ============================================================
# 5. RESULTS
# ============================================================
add_heading("5. Sonuçlar", level=1)

add_heading("5.1 Nihai Performans Karşılaştırması", level=2)
add_para(
    "Tablo 1, Ev 2 üzerinde persistence temel modeli ile üç adayin test seti metriklerini "
    "raporlar. Başlık sonuç: ARIMA(2, 1, 2) R2 = 0.876 ile kazanan, ancak önemsiz persistence "
    "temel modelini (R2 = 0.873) yenmesi ucuncu ondalık basamakta. LSTM (R2 = 0.748) ve "
    "FB-Prophet (R2 = 0.015) ikisi de persistence in altında kalır. Sonucu, Bölüm 5.5 te "
    "Diebold-Mariano testi ile ayrıntılı olarak inceleyecegiz, ama sezgisel çıkarım su: günlük "
    "konut enerjisindeki gecikme-1 otokorelasyonu butun ısı yapıyor; ARIMA bunu temiz biçimde "
    "yakalarken, LSTM ve Prophet bu yapıyı kuramiyor."
)

add_table_caption("Tablo 1. Ev 2 günlük tahminleri için persistence temel modeli (y_hat[t] = "
                  "y[t-1]) dahil dört yaklaşımın test seti performansi.")

table = doc.add_table(rows=5, cols=5, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['Model', 'MSE', 'RMSE (W)', 'MAE (W)', 'R2']
data = [
    ['Persistence',    '2,426.76', '49.26', '33.51', '0.873'],
    ['ARIMA(2,1,2)',   '2,370.01', '48.68', '32.75', '0.876'],
    ['LSTM (3-layer)', '4,821.35', '69.44', '52.25', '0.748'],
    ['FB-Prophet',     '18,851.52', '137.30', '111.70', '0.015'],
]
for i, h in enumerate(headers):
    c = table.rows[0].cells[i]
    c.text = h
    for p in c.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(10)
            r.font.name = 'Calibri'
for r_idx, row in enumerate(data):
    for c_idx, v in enumerate(row):
        c = table.rows[r_idx + 1].cells[c_idx]
        c.text = v
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.size = Pt(10)
                r.font.name = 'Calibri'

doc.add_paragraph()

add_figure("metrics_comparison.png",
           "Şekil 6. Üç model arasında dört değerlendirme metriginin yan-yana karşılaştırması. "
           "MSE, RMSE ve MAE için düşük olan iyi, R2 için yüksek olan iyi.",
           width=Inches(6.3))

add_heading("5.2 Tahmin Edilen ile Gercek Seri Karşılaştırması", level=2)
add_para(
    "Şekil 7, her modelin tahminlerini ayrılan test degerleriyle birlikte cizer. ARIMA panelinde yuvarlanan yeniden uydurma sayesinde model en son gozlenen degere sabitlendiginden gercek seri yakindan takip edilir. LSTM paneli trendi genel anlamda yakalar, ancak tepe ve diplerin keskinligini yumuşatır. Prophet paneli aslında parcali dogrusal bir taban çizgisi artırı mevsimsellik üretir; bu, ortalama olarak mantıklı kalsa da gercek verideki keskin hareketleri izlemekte başarısız olur."
)

add_figure("model_comparison.png",
           "Şekil 7. ARIMA, LSTM ve FB-Prophet için 122 günlük test penceresi boyunca gercek "
           "ile tahmin edilen günlük ortalama güç kiyaslamasi.",
           width=Inches(6.3))

add_heading("5.3 Eğitim Tanilari", level=2)
add_para(
    "Şekil 8 deki LSTM eğitim kaybi egrisi, ilk hizli yakinsamayi ve yaklaşık 12 epoch sonra platonun oluştuğunu gösterir; bu noktada erken durdurma tetiklenir. Doğrulama kaybi eğitim kaybini yakindan izler; bu, ana sorunun aşırı uyum olmadığını, modelin yumuşatılmış günlük sinyalden çıkarabileceği sinira ulastigini gösterir."
)

add_figure("lstm_training_history.png",
           "Şekil 8. LSTM eğitim ve doğrulama MSE kaybinin epoch lara göre degisimi.",
           width=Inches(5.0))

add_para(
    "Şekil 9 daki Prophet ayrıştırması belirgin bir yıllık trendi (en yüksek tüketim kisin, en düşük tüketim yazin) ve orta seviyede haftalık dongu sergiler. Bu bileşenler mantıklı olup keşifsel analizle örtüşür, ancak genel uyum kısa ufuklu hareketler yerine uzun ufuklu trendin etkisi altında kalır."
)

add_figure("prophet_components.png",
           "Şekil 9. Ev 2 nin eğitim bölümüne uydurulmuş FB-Prophet ayrıştırmasında altta yatan "
           "trend, haftalık mevsimsellik ve yıllık mevsimsellik bileşenleri.",
           width=Inches(5.5))

add_heading("5.4 Geliştirme İterasyonları", level=2)
add_para(
    "Tablo 1 deki nihai sonuçlar ilk denemede elde edilmedi. En iyi modelin R2 sini negatif bolgeden 0.876 ya taşıyan tasarim kararlarini belgeledik. Şekil 10 üç iterasyonu özetler."
)

add_figure("iteration_timeline.png",
           "Şekil 10. Tasarimin üç iterasyonu ve bunlarin test R2 sine etkisi. En çok fark "
           "yaratan değişiklik, tek seferlik çok-adımlı ARIMA tahmininden yuvarlanan "
           "tek-adım-önde protokole gecis oldu.",
           width=Inches(6.3))

doc.add_page_break()

# ============================================================
# 5.5 STATISTICAL SIGNIFICANCE (TR)
# ============================================================
add_heading("5.5 İstatistiksel Anlamlılık: Diebold-Mariano Testi", level=2)
add_para(
    "Ev 2 üzerinde yapılan ikili Diebold-Mariano testleri Tablo 2 deki sonuçları verir. "
    "Üç bulgu one cikiyor. Birincisi, ARIMA nin Persistence üzerindeki gozle gorulen üstünlüğü "
    "geleneksel hiçbir seviyede istatistiksel olarak anlamlı değil (p = 0.665). İki model bu test "
    "setinde pratik olarak ayırt edilemeyen tahminler üretiyor. İkincisi, LSTM ve Prophet "
    "Persistence den anlamlı şekilde daha kotu (her ikisinde de p < 0.001): model karmasikligini "
    "artırmak burada doğruluğu aktif olarak zedeliyor. Üçüncüsü, ARIMA LSTM ve Prophet i anlamlı "
    "olarak yener (her ikisinde p < 0.001); yanı secenek LSTM veya Prophet ise ARIMA mantıklı bir "
    "tercih, ancak tek satırlık bir gecikme-1 kurali da aynı ısı yapar."
)

add_table_caption("Tablo 2. Ev 2 için Harvey-Leybourne-Newbold küçük orneklem düzeltmesiyle "
                  "Diebold-Mariano test sonuçları. n = 123 test günü, iki tarafli p-değerleri.")

table2 = doc.add_table(rows=6, cols=4, style='Light Grid Accent 1')
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
dm_headers = ['Karşılaştırma (Model 1 vs Model 2)', 'DM Istatistigi', 'p-degeri', 'Sonuç']
dm_data = [
    ['ARIMA vs Persistence',  '-0.433',  '0.665',     'Anlamlı fark yok'],
    ['LSTM vs Persistence',   '+3.936',  '< 0.001',   'Persistence anlamlı olarak iyi'],
    ['Prophet vs Persistence','+7.949',  '< 0.001',   'Persistence anlamlı olarak iyi'],
    ['ARIMA vs LSTM',         '-3.765',  '< 0.001',   'ARIMA anlamlı olarak iyi'],
    ['ARIMA vs Prophet',      '-7.957',  '< 0.001',   'ARIMA anlamlı olarak iyi'],
]
for i, h in enumerate(dm_headers):
    c = table2.rows[0].cells[i]
    c.text = h
    for p in c.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(10)
            r.font.name = 'Calibri'
for r_idx, row in enumerate(dm_data):
    for c_idx, v in enumerate(row):
        c = table2.rows[r_idx + 1].cells[c_idx]
        c.text = v
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.size = Pt(10)
                r.font.name = 'Calibri'

doc.add_paragraph()

add_figure("dm_test_plot.png",
           "Şekil 11. Ev 2 için Diebold-Mariano istatistikleri. Yesil çubuklar yüzde 5 seviyesinde "
           "anlamlı (|DM| > 1.96), gri çubuklar anlamlı değil. ARIMA ile Persistence "
           "karşılaştırması anlamlı olmayan tek ciftir.",
           width=Inches(6.0))

# ============================================================
# 5.6 CROSS-HOUSE VALIDATION (TR)
# ============================================================
add_heading("5.6 Hane-Aşırı Doğrulama", level=2)
add_para(
    "Ev 2 deki sıra lamasinin tek bir hanenin tesadufu olmadığını gostermek için tüm boru hattini "
    "Ev 1 ve Ev 5 te tekrarladik. Bu iki ev farklı kullanım duzeni ve cihaz karisimina sahip "
    "(Ev 1 de dondurucular ve kurutucu, Ev 5 te belirgin şekilde daha yüksek günlük ortalama güç "
    "bulunuyor); dolayısıyla üç evde de ayakta kalan herhangi bir sıralama tesadufi olamaz. "
    "Tablo 3, hane başına R2 degerlerini sunar."
)

add_table_caption("Tablo 3. Üç REFIT hanesinde model başına R2. Persistence önemsiz temel "
                  "modeldir; ARIMA her evde dar bir farkla kazanır, LSTM ve Prophet ise tutarlı "
                  "şekilde Persistence in altında kalır.")

table3 = doc.add_table(rows=5, cols=5, style='Light Grid Accent 1')
table3.alignment = WD_TABLE_ALIGNMENT.CENTER
ch_headers = ['Model', 'Ev 1', 'Ev 2', 'Ev 5', 'Yargı']
ch_data = [
    ['Persistence', '0.963', '0.873', '0.901', 'Güçlü temel model'],
    ['ARIMA',       '0.976', '0.876', '0.905', 'Her evde en iyi'],
    ['LSTM',        '0.859', '0.748', '0.647', 'Her evde Persistence altında'],
    ['FB-Prophet',  '-2.315', '0.015', '-1.186', 'Sik sik negatif R2'],
]
for i, h in enumerate(ch_headers):
    c = table3.rows[0].cells[i]
    c.text = h
    for p in c.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(10)
            r.font.name = 'Calibri'
for r_idx, row in enumerate(ch_data):
    for c_idx, v in enumerate(row):
        c = table3.rows[r_idx + 1].cells[c_idx]
        c.text = v
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.size = Pt(10)
                r.font.name = 'Calibri'

doc.add_paragraph()

add_figure("cross_house_comparison.png",
           "Şekil 12. Hane başına model R2 si. ARIMA ile Persistence arasındaki fark üç evde de "
           "küçük, LSTM tutarlı şekilde Persistence in altında, Prophet Ev 1 ve Ev 5 te negatif "
           "R2 üretiyor.",
           width=Inches(6.3))

add_para(
    "İki bulgu burada vurgulanmalı. Birincisi, ARIMA ile Persistence arasındaki fark her evde "
    "küçük; yanı Ev 2 sonucu sıra dışı değil. İkincisi, Prophet üç evden ikisinde negatif R2 "
    "ureterek koşulsuz ortalamadan bile daha kotu tahmin ediyor; bu, tek hane bulgusundan daha "
    "güçlü bir ifade."
)

# ============================================================
# 5.7 PRACTICAL INTERPRETATION (TR)
# ============================================================
add_heading("5.7 Hata Degerlerinin Pratik Yorumu", level=2)
add_para(
    "RMSE değerleri, alan birimlerine çevrildiğinde daha kolay anlaşılır. Ev 2 için ortalama test "
    "tüketimi 434.8 W, yanı ARIMA nin 48.7 W luk RMSE si tipik günlük seviyenin yaklaşık yüzde "
    "11.2 sine denk gelir. Enerji cinsine çevrildiğinde, 24 saat boyunca sürdürülen 48.7 W luk "
    "ortalama-güç hatası yaklaşık 1.17 kWh/gün eder. Birleşik Krallık 2025 konut elektrik birim "
    "fiyati yaklaşık 0.27 GBP/kWh olduğunda, bu tek hane için yaklaşık 0.32 GBP/gün beklenen "
    "faturalama hatası anlamına gelir; aylık yaklaşık 10 GBP. Toplulastirilmis bir bilanco "
    "ortaminda gün-önceki dengesizlik 0.10 GBP/kWh fiyatlandirildiginda maliyet yaklaşık "
    "0.12 GBP/gün e duser. Bu sayilar hane başına küçük olsa da portföy büyüklüğü ile dogrusal "
    "olarak ölçeklenir."
)

# ============================================================
# 6. DISCUSSION
# ============================================================
add_heading("6. Tartışma", level=1)

add_heading("6.1 Persistence Sonucu ve Bize Söylediği", level=2)
add_para(
    "Bu çalışmanın en önemli bulgusu ARIMA nin diğer modelleri yenmesi değil, ARIMA nin persistence i "
    "ancak duz duz yenmesi ve farkın istatistiksel olarak anlamlı olmamasidir. Persistence yarinin "
    "bugune eşit oldugunu tahmin eder: hiçbir ogrenilmis parametresi olmayan tek satırlık bir kural. "
    "Ev 2 de R2 = 0.873 e ulasiyor, ARIMA nin 0.876 sindan sadece 0.003 düşük. Diebold-Mariano testi "
    "(p = 0.665) iki modelin karesel hata kayiplarinin ayırt edilemez oldugunu doğrular. Hane-aşırı "
    "sonuçlar aynı resmi cizer: ARIMA nin Persistence üzerindeki R2 farkı Ev 1 de 0.013, Ev 5 te 0.004."
)
add_para(
    "Mekanik neden, 7 günlük yumuşatılmış günlük tüketiminin güçlü gecikme-1 otokorelasyonudur. "
    "Dun bugunu mükemmel şekilde tahmin eder ve dunu doğrudan kullanan herhangi bir model neredeyse "
    "aynı performansi gösterir. ARIMA(2, 1, 2) iki gecikme daha ve iki hareketli-ortalama terimi "
    "ekler, ancak ek bilgi gecikme-1 e kiyasla kucuktur. Pratik sonuç net: bu granulariteyde ve "
    "bu sinyal üzerinde, tek satırlık bir kural zaten oraddatahmin edilebilir yapıyı yakaliyorken "
    "ağır makineler kullanmayın."
)

add_heading("6.2 LSTM Neden Önemsiz Temel Modelden de Düşük Performans Gösterdi", level=2)
add_para(
    "LSTM yalnizca ARIMA ya kaybetmiyor, Persistence e karşı da anlamlı şekilde kaybediyor (Ev 2 de "
    "p < 0.001, diğer iki evde benzer örüntü). Bu daha ağır bir sonuç: 128.929 parametreli sınır ağı, "
    "yarinin bugune eşit olacagini tahmin etmekten daha kotu. Eğitim seti 490 günlük gözlemden oluşur "
    "ve 14 günlük geriye bakış ile yalnizca 476 girdi dizisi verir. Bu kadar az veriyle ag aşırı "
    "kapasiteye sahip; yumuşak ortalama davranisi ogreniyor ve Persistence in tanim geregi yakaladığı "
    "günden güne hareketleri aşırı yumuşatıyor. Son dönem çalışmaları (Gasparin ve diğerleri, 2024) "
    "benzer örüntüyü belgeler: altı aydan kısa eğitim verisinde basit temel modeller derin modelleri "
    "geride bırakır. Bizim verimiz daha uzun, ancak yumuşatılmış günlük sinyal hala derin ogrenmenin "
    "net olarak negatif olduğu bolgede gorunuyor."
)

add_heading("6.3 Prophet Neden En Düşük Skoru Aldi", level=2)
add_para(
    "Prophet in Ev 2 deki sifir civarındaki R2 si zaten cesaret kirici degildi. Hane-aşırı doğrulama "
    "resmi daha da kotulestiriyor: Ev 1 de R2 = -2.315 ve Ev 5 te R2 = -1.186, yanı Prophet koşulsuz "
    "eğitim ortalamasindan önemli olcude daha kotu tahmin ediyor. İki etken devrededir. Birincisi, "
    "Prophet tüm test ufkunu test seti gercek degerlerine erismeden tek seferde tahmin eder; bu, "
    "yuvarlanan Persistence, ARIMA ve LSTM kurulumlarina kiyasla onu yapisal olarak dezavantajli "
    "bırakır. İkincisi, Prophet in güçlü yanı yavas mevsimsel trendleri ve takvim etkilerini "
    "yakalamaktir; yumuşatılmış günlük konut verisindeki baskın sinyal bir önceki gün olup Prophet in "
    "ayrıştırmasının hedeflemedigi bileşendir. Adil bir karşılaştırma Prophet i yuvarlanan yeniden "
    "uydurma modunda çalıştırmayı gerektirir, ancak veri karakteriyle yapisal uyumsuzluk yine kalır."
)

add_heading("6.4 Kisitlamalar ve Geçerlilik Tehditleri", level=2)
add_para(
    "Üç kısıtlama sonuclarimizi nitelendirir."
)
add_bullet("Yirmi değil, üç hane. Hane-aşırı doğrulama orijinal tek-hane bulgusunu guclendirir, "
           "ancak kalan 17 REFIT evini henuz test etmiyor.")
add_bullet("Tek degiskenli kurulum. Modellerin hicbiri hava, takvim veya cihaz düzeyinde dışsal "
           "veri almaz. Bunlarin eklenmesi tabloyu, özellikle daha zengin girdileri emmek için "
           "tasarlanan LSTM lehine, değiştirebilir.")
add_bullet("Ağır yumuşatma. 7 günlük hareketli ortalama yüksek frekansli içeriği ortadan kaldirir. "
           "Günlük rampa bilgisine ihtiyaç duyan bir operasyonel tahminci için bu seçim yeniden "
           "gozden geçirilmelidir; persistence avantaji daha ince zaman ölçeklerinde kuculebilir.")

add_heading("6.5 Referans Literatür ile Karşılaştırma", level=2)
add_para(
    "Burada gozlemlenen ARIMA, LSTM ve Prophet siralamasi, Buluc, Sevli ve Yunlu (2025) un İstanbul güneş üretim verisinde benzer yuvarlanan protokol altında raporladigi (R2 = 0.97) ARIMA performansiyla örtüşür. Demirtop ve Sevli (2024) un LSTM in ARIMA yi gectigi ruzgar hizi çalışmasından ise ayrışır; en olasi neden o çalışmada çok yıllık saatlik verinin kullanilmis olmasidir. CNN-LSTM-Transformer (Limouni ve diğerleri, 2023) ve STL-Prophet-LSTM (Xie ve diğerleri, 2022) gibi hibrit modeller bileşenlerinden tutarlı şekilde daha iyi sonuç verir ve burada karsilastirilan tek-model temellerinin doğal sonraki adımını oluşturur."
)

# ============================================================
# 7. CONCLUSION
# ============================================================
add_heading("7. Sonuç ve Gelecek Çalışmalar", level=1)

add_para(
    "Persistence temel modeli dahil dört yöntem kullanilarak üç Birleşik Krallık hanesinin günlük "
    "ortalama gucunu tahmin eden tekrar üretilebilir bir boru hatti kurduk. Yuvarlanan tek-adım-önde "
    "protokol ve sabit rastgele tohum altında, ARIMA(2, 1, 2) her evde en iyi model oldu, ancak tek "
    "satırlık persistence kurali üzerindeki marji istatistiksel olarak anlamlı degildi (Ev 2 de "
    "Diebold-Mariano p = 0.665 ve Ev 1 ile Ev 5 te benzer küçük farklar). LSTM ve FB-Prophet üç "
    "ortamin tümünde persistence e anlamlı olarak kaybetti; Prophet üç evden ikisinde negatif R2 "
    "uretti. Cikardigimiz ders, bu granulariteyde konut günlük enerjisinin gecikme-1 otokorelasyonu "
    "tarafından domine edildiği ve önemsiz temel modelin üzerine model kapasitesi eklemenin olculebilir "
    "kazanç saglamadigi (ARIMA) veya aktif olarak zararli olduğu (LSTM, Prophet) yonundedir."
)

add_para(
    "Metodolojik ders de en az aynı kadar önemli: yalnizca karmasik modelleri birbirine karşı raporlamak, "
    "tüm model sinifinin gereksiz olduğu gercegini gizleyebilir. Persistence temel modeli ile Diebold-"
    "Mariano anlamlılık testi birlikte, temiz bir ARIMA zaferi gibi gorunen sonucu daha durust 'veri "
    "modellerin önemli olmasi için fazla tahmin edilebilir' sonucuna donusturur."
)

add_para(
    "İleride yapilacak çalışmalar için dört yön one cikiyor."
)
add_bullet("Gecikme-1 otokorelasyonunun daha zayıf olduğu ve derin modellerin gercekten katkı "
           "sağlayabileceği daha yüksek frekansli hedefte (saatlik veya 15 dakikalık) analizi tekrarlamak.")
add_bullet("Hava (sıcaklık, ışınım) ve takvim ozelliklerini SARIMAX ve çok degiskenli LSTM "
           "aracılığıyla eklemek; bu daha ağır modellere persistence in kullanamayacagi bilgiyi verir.")
add_bullet("Hane-aşırı karsilastirmayi ucten tüm 20 REFIT evine uzatmak ve hane-aşırı transfer "
           "ogrenmesini keşfetmek.")
add_bullet("Ayrıştırma ve sinirsel dizi ogrenmenin güçlü yanlarini birleştiren STL-Prophet-LSTM ve "
           "CNN-LSTM-Transformer gibi hibrit modeller kurmak.")

add_heading("Tekrarlanabilirlik Notu", level=2)
add_para(
    "Tüm deneyler NumPy, Python in random modulu ve TensorFlow için sabit rastgele tohum (42) kullanır. Veri seti, REFIT in herkese açık Kaggle aynasından indirilir ve ön işleme, eğitim ve değerlendirme tek bir Python betiğinde (energy_forecasting.py) kodlanir. Bu rapordaki sonuçlar Python 3.13, TensorFlow 2.x, statsmodels 0.14 ve Prophet 1.3 üzerinde elde edilmiştir."
)

# ============================================================
# REFERENCES
# ============================================================
add_heading("Kaynakca", level=1)

refs = [
    "Atalay, B. A., and Zor, K. (2025). XGBoost ile hidroelektrik enerji tahmini. "
    "Çukurova Üniversitesi Mühendislik Fakültesi Dergisi, 40(1), 205-218.",

    "Berus, Y., and Yakut, Y. B. (2024). Derin öğrenme (1D-CNN, RNN, LSTM, BiLSTM) ile enerji "
    "tüketim tahmini: Diyarbakır AVM örneği. Dicle University Journal of Engineering, 15(2), "
    "311-322.",

    "Bülüç, M., Sevli, O., and Yünlü, L. (2025). Time Series Analysis of Solar Energy Production "
    "Based on Weather Conditions. GU J Sci, Part A, 12(4), 1060-1077.",

    "Çolak, M. B., and Özhan, E. (2025). Renewable energy forecasting in Turkey: Analytical "
    "approaches. Journal of Intelligent Systems: Theory and Applications, 8(1), 25-34.",

    "Demirtop, A., and Sevli, O. (2024). Wind speed prediction using LSTM and ARIMA time series "
    "analysis models: A case study of Gelibolu. Turkish Journal of Engineering, 8(3), 524-536.",

    "Gasparin, A., Lukovic, S., and Alippi, C. (2024). Load Forecasting for Households and "
    "Energy Communities: Are Deep Learning Models Worth the Effort? arXiv:2501.05000.",

    "Limouni, T., Yaagoubi, R., Bouziane, K., Guissi, K., and Baali, E. H. (2023). Solar Energy "
    "Production Forecasting Based on a Hybrid CNN-LSTM-Transformer Model. Mathematics, 11(3), 676.",

    "Murray, D., Stankovic, L., and Stankovic, V. (2017). An electrical load measurements "
    "dataset of United Kingdom households from a two-year longitudinal study. Scientific Data, "
    "4, 160122.",

    "Rahman, M. R., and others (2025). Time-series and deep learning approaches for renewable "
    "energy forecasting in Dhaka: a comparative study of ARIMA, SARIMA, and LSTM models. "
    "Discover Sustainability, 6, 01733.",

    "Sakib, N., and others (2024). Deep learning-driven hybrid model for short-term load "
    "forecasting and smart grid information management. Scientific Reports, 14, 63262.",

    "Tamay, M., and Türker, G. F. (2024). Machine learning based energy forecasting for "
    "photovoltaic solar plants. Yekarum, 9(2), 128-146.",

    "Triebe, O., Hewamalage, H., Pilyugina, P., Laptev, N., Bergmeir, C., and Rajagopal, R. "
    "(2021). NeuralProphet: Explainable Forecasting at Scale. arXiv:2111.15397.",

    "Vargas-Forero, V. M., Manotas-Duque, D. F., and Trujillo, L. (2025). Comparative study of "
    "forecasting methods to predict the energy demand for the market of Colombia. International "
    "Journal of Energy Economics and Policy, 15(1), 65-76.",

    "Xie, J., Li, Z., Zhou, Z., and Liu, S. (2022). A hybrid forecasting model using LSTM and "
    "Prophet for energy consumption with decomposition of time series data. PeerJ Computer "
    "Science, 8, e1001.",
]

for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.first_line_indent = Cm(-0.6)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f"[{i}] {ref}")
    r.font.size = Pt(9.5)
    r.font.name = 'Calibri'


# ============================================================
# SAVE DOCX
# ============================================================
doc.save(OUTPUT_DOCX)
print(f"DOCX saved: {OUTPUT_DOCX}")
print(f"DOCX size: {os.path.getsize(OUTPUT_DOCX)/1024:.0f} KB")

# ============================================================
# CONVERT TO PDF
# ============================================================
print("\nConverting to PDF...")
try:
    from docx2pdf import convert
    convert(OUTPUT_DOCX, OUTPUT_PDF)
    print(f"PDF saved: {OUTPUT_PDF}")
    print(f"PDF size: {os.path.getsize(OUTPUT_PDF)/1024:.0f} KB")
except Exception as e:
    print(f"docx2pdf failed: {e}")
    print("Will try alternative PDF conversion methods.")
