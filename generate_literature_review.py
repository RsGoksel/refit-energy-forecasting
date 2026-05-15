"""
Generate Literature Review DOCX for the project:
"Household Energy Consumption Forecasting Using Time Series Models"

Structured per the instructor's guide:
- Introduction
- Thematic Review
- Critical Analysis
- Conclusion
- Summary comparison table
- References (10+ articles)
"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE_DIR = os.path.dirname(__file__)
OUTPUT_PATH = os.path.join(BASE_DIR, "EBT629E_Literature_Review.docx")

doc = Document()

# --- Page margins ---
for section in doc.sections:
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# --- Default style: Calibri 11pt, 1.3 line spacing ---
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(5)
style.paragraph_format.line_spacing = 1.3


def add_heading_styled(text, level=1, color=RGBColor(0, 51, 102)):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
        run.font.name = 'Calibri'
    return h


def add_para(text, bold_lead=None, justify=True):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        r.font.name = 'Calibri'
        r.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.name = 'Calibri'
    r2.font.size = Pt(11)
    return p


# ============================================================
# TITLE
# ============================================================
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_after = Pt(2)
run = title_p.add_run("Literature Review")
run.bold = True
run.font.size = Pt(16)
run.font.name = 'Calibri'
run.font.color.rgb = RGBColor(0, 51, 102)

subtitle_p = doc.add_paragraph()
subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_p.paragraph_format.space_after = Pt(2)
run = subtitle_p.add_run("Household Energy Consumption Forecasting Using Time Series Models")
run.bold = True
run.font.size = Pt(13)
run.font.name = 'Calibri'

course_p = doc.add_paragraph()
course_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
course_p.paragraph_format.space_after = Pt(4)
run = course_p.add_run("EBT 629E \u2013 Artificial Intelligence | Istanbul Technical University")
run.font.size = Pt(10)
run.font.name = 'Calibri'
run.font.color.rgb = RGBColor(80, 80, 80)

info_p = doc.add_paragraph()
info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
info_p.paragraph_format.space_after = Pt(10)
run = info_p.add_run(
    "Nurullah Y\u0131ld\u0131r\u0131m (301252004) | "
    "Kadir G\u00f6ksel G\u00fcnd\u00fcz (301241077) | "
    "Furkan \u00c7\u0131nar (301212001)   \u2013   April 20, 2026"
)
run.font.size = Pt(10)
run.font.name = 'Calibri'
run.font.color.rgb = RGBColor(80, 80, 80)

# ============================================================
# 1. INTRODUCTION
# ============================================================
add_heading_styled("1. Introduction", level=2)

add_para(
    "Accurate forecasting of household energy consumption has become a cornerstone of smart grid "
    "operation, demand-side management, and the integration of distributed energy resources. "
    "With the widespread deployment of smart meters and the rapid growth of machine learning (ML) "
    "techniques, time series forecasting at the residential level has received increasing attention "
    "over the last five years. Residential consumption is, however, highly stochastic and is driven "
    "by occupant behavior, appliance diversity, weather conditions, and socio-economic factors, which "
    "makes it a particularly challenging forecasting problem. This review focuses exclusively on "
    "peer-reviewed and conference publications from 2020 onwards that apply time series or hybrid ML "
    "models (primarily ARIMA, LSTM, Facebook Prophet, and their extensions) to energy-systems "
    "forecasting tasks, spanning residential load, renewable generation (solar, wind, hydro), and "
    "grid-level electricity demand. The goal is to identify dominant methodological trends, compare "
    "reported accuracies, highlight remaining research gaps, and situate the present project within "
    "this landscape."
)

# ============================================================
# 2. THEMATIC REVIEW
# ============================================================
add_heading_styled("2. Thematic Review", level=2)

add_heading_styled("2.1 Classical Statistical Models (ARIMA and SARIMA Family)", level=3)
add_para(
    "ARIMA and its variants remain competitive baselines for energy forecasting, particularly for "
    "short-term and stationary series. B\u00fcl\u00fc\u00e7 and others (2025) reported that ARIMA achieved the lowest "
    "error rates (R2 = 0.97) in a three-way comparison with LSTM and FB-Prophet on data from the "
    "\u0130kitelli Solar Power Plant in Istanbul, arguing that its strength lies in handling short, "
    "autocorrelated sequences. Tamay and T\u00fcrker (2024), in contrast, found that tree-based methods "
    "(KNN, RMSE = 10.6) outperformed ARIMA on solar PV data from Denizli, indicating that ARIMA\u2019s "
    "advantage diminishes when nonlinear exogenous variables dominate. Rahman and others (2025) compared "
    "ARIMA, SARIMA, and LSTM for renewable-energy generation in Dhaka and confirmed that classical "
    "models handle linear trends and seasonality well but are consistently outperformed by LSTM once "
    "long-range nonlinear dependencies appear. The overall picture from these works is that ARIMA "
    "stays a strong benchmark when the series is short, stationary after differencing, and dominated "
    "by linear autocorrelation, but its performance plateaus once external drivers or long-range "
    "dependencies become important."
)

add_heading_styled("2.2 Deep Learning Approaches (LSTM, BiLSTM, CNN-LSTM)", level=3)
add_para(
    "Recurrent neural networks, especially LSTM, are the most studied deep learning architecture for "
    "residential load forecasting. Demirtop and Sevli (2024) showed that LSTM outperformed ARIMA on "
    "wind speed data from Gelibolu (RMSE = 3.10 vs. higher for ARIMA), while Berus and Yakut (2024) "
    "compared 1D-CNN, RNN, LSTM, and BiLSTM on shopping-mall consumption data and reported that "
    "BiLSTM achieved the best performance (RMSE = 0.050, R2 = 0.93) thanks to its bidirectional "
    "context. \u00c7olak and \u00d6zhan (2025) extended this comparison to a macro-level renewable energy "
    "series for Turkey (1960\u20132025) and again found LSTM superior to ARIMA(1,2,1) and NNAR. Focusing "
    "specifically on the REFIT dataset, recent work (ResearchSquare, 2025) confirmed that LSTM captures "
    "complex temporal patterns in 8-second smart-meter data more effectively than ARIMA, at the cost "
    "of substantially higher computational requirements (up to 2 hours of GPU training and ~4.7 GB "
    "memory). Across these works, LSTM gains most from long, multivariate, "
    "and nonlinear series, but its advantage shrinks on short or heavily smoothed data. This finding is "
    "echoed by Gasparin and others (2024), who showed that simple persistence baselines can "
    "outperform deep models when training data is limited to six months or less."
)

add_heading_styled("2.3 Automated Decomposition Models (FB-Prophet)", level=3)
add_para(
    "Facebook Prophet is widely adopted because of its human-in-the-loop design and its automatic "
    "handling of trend, seasonality, and holiday effects. Vargas-Forero and others (2025) showed that "
    "Prophet slightly outperformed ARIMA and LSTM for 24-hour electricity demand forecasting in "
    "Colombia (RMSE = 428.5 MW, MAPE = 4.1%), particularly around holidays. On the other hand, "
    "B\u00fcl\u00fc\u00e7 and others (2025) reported Prophet as the weakest of the three tested models on short, "
    "high-variance solar series. Xie and others (2022) proposed an STL-decomposed Prophet-LSTM hybrid for "
    "building energy consumption, where Prophet handled trend and seasonality while LSTM modeled the "
    "residuals, achieving lower RMSE than either model in isolation. A Neural-Prophet variant "
    "(Triebe and others, 2021) has been proposed to mitigate Prophet\u2019s weaknesses by combining its "
    "interpretability with autoregressive neural components. Overall, Prophet does best when "
    "the signal has strong weekly or yearly seasonality and when calendar effects (weekends, holidays) "
    "materially drive demand, while it struggles with the abrupt fluctuations typical of single-"
    "household consumption."
)

add_heading_styled("2.4 Hybrid and Ensemble Models for Energy Forecasting", level=3)
add_para(
    "Hybrid architectures that combine statistical and deep components consistently appear at the top "
    "of reported energy-forecasting benchmarks. Limouni and others (2023) proposed a CNN-LSTM-Transformer "
    "hybrid for solar energy production, where the CNN extracted local spatial features, LSTM captured "
    "mid-range temporal dependencies, and Transformer attention modeled long-range interactions; the "
    "integrated model outperformed ARIMA, DeepAR, and Prophet baselines on PV datasets. Xie and others (2022) "
    "similarly demonstrated that an STL-Prophet-LSTM decomposition pipeline improved energy-consumption "
    "forecasting accuracy over any single component. Sakib and others (2024) combined CNN feature extractors "
    "with LSTM decoders for multi-step residential load prediction on the smart-grid dataset, again "
    "outperforming pure recurrent models. Atalay and Zor (2025) took a different route, using XGBoost "
    "with 1,000 trees for hydroelectric generation (R2 = 0.96, RMSE = 7.29), confirming that gradient-"
    "boosted ensembles are highly competitive when sufficient engineered features are available. These "
    "results reinforce the observation that no single model dominates across all operational regimes "
    "and that targeted decomposition often yields the strongest empirical performance in energy systems."
)

add_heading_styled("2.5 Transformer-Based and Emerging Architectures", level=3)
add_para(
    "Transformer-based models are the most recent addition to the forecasting toolbox. Gaviria-Chav\u00e9z "
    "and others (2025) compared SARIMA through Transformer variants on the REFIT dataset and reported that "
    "the Temporal Fusion Transformer (TFT) achieved the best point-forecast accuracy (RMSE = 481.94), "
    "while LSTM provided the best-calibrated probabilistic forecasts. Li and others (2024) extended this "
    "with a spatiotemporal graph-attention Transformer for multivariate residential load forecasting, "
    "reporting at least 14.7% improvement over deep-learning benchmarks in multi-step settings. However, "
    "Transformers introduce considerable complexity and data requirements, which has motivated research "
    "into transfer learning across households (Li and others, 2025), reducing MAE by roughly 16% for 24-hour "
    "forecasts when source datasets are combined."
)

# ============================================================
# 3. SUMMARY TABLE
# ============================================================
add_heading_styled("3. Comparative Summary of Reviewed Studies", level=2)

cap = doc.add_paragraph()
cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
cap.paragraph_format.space_after = Pt(2)
run = cap.add_run("Table 1. Summary of key reviewed studies, methods, datasets, and headline results.")
run.bold = True
run.font.size = Pt(9)
run.font.name = 'Calibri'

table = doc.add_table(rows=1, cols=5, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['Study', 'Domain / Data', 'Methods Compared', 'Best Model (R2 / RMSE)', 'Gap / Limitation']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(9)
            r.font.name = 'Calibri'

rows_data = [
    ['B\u00fcl\u00fc\u00e7 and others (2025)', '\u0130kitelli PV, Istanbul', 'ARIMA, LSTM, FB-Prophet',
     'ARIMA (R2 = 0.97)', 'Short (1-month) dataset; no irradiance feature'],

    ['Demirtop & Sevli (2024)', 'Wind, Gelibolu', 'ARIMA, LSTM',
     'LSTM (RMSE = 3.10)', 'Single site; no hybrid models'],

    ['Berus & Yakut (2024)', 'Mall, Diyarbak\u0131r', '1D-CNN, RNN, LSTM, BiLSTM',
     'BiLSTM (R2 = 0.93)', 'Aggregate series only'],

    ['\u00c7olak & \u00d6zhan (2025)', 'Turkey renewable', 'LSTM, NNAR, ARIMA, ELM',
     'LSTM (long-term best)', 'Annual resolution; coarse'],

    ['Tamay & T\u00fcrker (2024)', 'PV, Denizli', 'XGBoost, KNN, RF, GBM',
     'KNN (R2 = 0.97)', 'Tabular features only'],

    ['Atalay & Zor (2025)', 'Hydro, Aslanta\u015f', 'XGBoost',
     'XGBoost (R2 = 0.96)', 'Single model baseline'],

    ['Vargas-Forero and others (2025)', 'Colombia demand', 'ARIMA, LSTM, Prophet',
     'Prophet (MAPE = 4.1%)', 'No exogenous weather'],

    ['Limouni and others (2023)', 'Solar PV generation', 'CNN-LSTM-Transformer hybrid',
     'Hybrid > ARIMA/Prophet', 'High model complexity'],

    ['Xie and others (2022)', 'Building energy use', 'STL-Prophet-LSTM',
     'Hybrid (lowest RMSE)', 'Requires STL tuning'],

    ['Rahman and others (2025)', 'Renewables, Dhaka', 'ARIMA, SARIMA, LSTM',
     'LSTM (lowest RMSE)', 'Single country scope'],

    ['Gaviria-Ch\u00e1vez and others (2025)', 'REFIT households', 'SARIMA \u2192 TFT',
     'TFT (RMSE = 481.9)', 'High compute cost'],

    ['ResearchSquare (2025)', 'Smart meter (household)', 'ARIMA, LSTM',
     'LSTM (MAPE = 3.0%)', 'Two households only'],

    ['Sakib and others (2024)', 'Residential load', 'CNN-LSTM hybrid',
     'Hybrid > pure LSTM', 'Requires feature tuning'],
]

for row_data in rows_data:
    row = table.add_row()
    for i, val in enumerate(row_data):
        cell = row.cells[i]
        cell.text = val
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)
                r.font.name = 'Calibri'

doc.add_paragraph()

# ============================================================
# 4. CRITICAL ANALYSIS
# ============================================================
add_heading_styled("4. Critical Analysis", level=2)

add_para(
    "Three main patterns stand out in the reviewed literature. ",
    bold_lead="Trend 1 \u2013 Context-dependent model ranking. "
)
add_para(
    "No single architecture is universally best. ARIMA tends to win on short, stationary, linear energy "
    "series (B\u00fcl\u00fc\u00e7 and others, 2025; Rahman and others, 2025), LSTM/BiLSTM win on long and nonlinear signals "
    "(Demirtop & Sevli, 2024; Berus & Yakut, 2024), Prophet wins when calendar effects dominate demand "
    "(Vargas-Forero and others, 2025), and Transformers win on very long multivariate data at the cost of "
    "compute (Gaviria-Ch\u00e1vez and others, 2025).",
    bold_lead=""
)

add_para(
    "",
    bold_lead="Trend 2 \u2013 Hybridization yields the largest gains. "
)
add_para(
    "Decomposition-based hybrids such as CNN-LSTM-Transformer (Limouni and others, 2023) for solar PV, "
    "STL-Prophet-LSTM (Xie and others, 2022) for building energy, and CNN-LSTM (Sakib and others, 2024) for "
    "residential load consistently outperform their individual components, suggesting that explicitly "
    "separating trend, seasonality, and residual dynamics is a fruitful design principle in energy "
    "applications.",
    bold_lead=""
)

add_para(
    "",
    bold_lead="Trend 3 \u2013 Persistent data-scale and feature limitations. "
)
add_para(
    "Many studies report negative or unstable results when training windows are shorter than six months "
    "or when key exogenous drivers (direct irradiance, household occupancy, appliance metadata) are "
    "missing. Gasparin and others (2024) explicitly show that persistence baselines beat deep models under "
    "these conditions, highlighting an ongoing need for longer, richer datasets and transfer-learning "
    "strategies.",
    bold_lead=""
)

add_para(
    "",
    bold_lead="Research gaps. "
)
add_para(
    "First, relatively few studies apply all three of ARIMA, LSTM, and Prophet to genuinely "
    "household-level (not aggregated) data under a rolling-forecast protocol. Second, the trade-off "
    "between accuracy and computational cost is rarely quantified, even though it is decisive for "
    "edge-deployed smart meters. Third, publicly available datasets such as REFIT are still "
    "under-exploited for cross-household transfer learning.",
    bold_lead=""
)

# ============================================================
# 5. CONCLUSION
# ============================================================
add_heading_styled("5. Conclusion and Positioning of the Project", level=2)

add_para(
    "The reviewed literature indicates that no forecasting model is universally optimal; rather, the "
    "best choice depends on data horizon, noise level, and the presence of exogenous drivers. Building "
    "on this insight, our project will systematically compare ARIMA, LSTM, and FB-Prophet on the REFIT "
    "Smart Home dataset (Murray and others, 2017) using a single household as the test bed. In contrast to "
    "most prior work, we will adopt a rolling one-step-ahead protocol to better approximate real-time "
    "operation, and we will explicitly evaluate the accuracy-versus-complexity trade-off using MSE, "
    "RMSE, MAE, and R2. Early experiments already confirm the ARIMA advantage observed by "
    "B\u00fcl\u00fc\u00e7 and others (2025) on short horizons, motivating future work on hybrid decomposition models "
    "and transfer learning across the 20 REFIT households. The literature has so far explored these "
    "directions only in isolation."
)

# ============================================================
# 6. REFERENCES
# ============================================================
add_heading_styled("References", level=2)

refs = [
    "Atalay, B. A., & Zor, K. (2025). XGBoost ile hidroelektrik enerji tahmini. "
    "\u00c7ukurova \u00dcniversitesi M\u00fchendislik Fak\u00fcltesi Dergisi, 40(1), 205\u2013218.",

    "Berus, Y., & Yakut, Y. B. (2024). Derin \u00f6\u011frenme (1D-CNN, RNN, LSTM, BiLSTM) ile enerji "
    "t\u00fcketim tahmini: Diyarbak\u0131r AVM \u00f6rne\u011fi. Dicle University Journal of Engineering, 15(2), 311\u2013322.",

    "B\u00fcl\u00fc\u00e7, M., Sevli, O., & Y\u00fcnl\u00fc, L. (2025). Time Series Analysis of Solar Energy Production "
    "Based on Weather Conditions. GU J Sci, Part A, 12(4), 1060\u20131077.",

    "\u00c7olak, M. B., & \u00d6zhan, E. (2025). Renewable energy forecasting in Turkey: Analytical approaches. "
    "Journal of Intelligent Systems: Theory and Applications, 8(1), 25\u201334.",

    "Demirtop, A., & Sevli, O. (2024). Wind speed prediction using LSTM and ARIMA time series analysis "
    "models: A case study of Gelibolu. Turkish Journal of Engineering, 8(3), 524\u2013536.",

    "Gasparin, A., Lukovic, S., & Alippi, C. (2024). Load Forecasting for Households and Energy "
    "Communities: Are Deep Learning Models Worth the Effort? arXiv:2501.05000.",

    "Gaviria-Ch\u00e1vez, F., and others (2025). Robust Probabilistic Load Forecasting for a Single Household: "
    "A Comparative Study from SARIMA to Transformers on the REFIT Dataset. arXiv:2512.00856.",

    "Li, K., Huang, W., Hu, G., & Li, J. (2024). Enhancing multivariate, multi-step residential load "
    "forecasting with spatiotemporal graph attention-enabled Transformer. International Journal of "
    "Electrical Power & Energy Systems, 159, 110049.",

    "Limouni, T., Yaagoubi, R., Bouziane, K., Guissi, K., & Baali, E. H. (2023). Solar Energy Production "
    "Forecasting Based on a Hybrid CNN-LSTM-Transformer Model. Mathematics, 11(3), 676.",

    "Murray, D., Stankovic, L., & Stankovic, V. (2017). An electrical load measurements dataset of "
    "United Kingdom households from a two-year longitudinal study. Scientific Data, 4, 160122.",

    "Rahman, M. R., and others (2025). Time-series and deep learning approaches for renewable energy "
    "forecasting in Dhaka: a comparative study of ARIMA, SARIMA, and LSTM models. Discover "
    "Sustainability, 6, 01733.",

    "Sakib, N., and others (2024). Deep learning-driven hybrid model for short-term load forecasting and "
    "smart grid information management. Scientific Reports, 14, 63262.",

    "Tamay, M., & T\u00fcrker, G. F. (2024). Machine learning based energy forecasting for photovoltaic "
    "solar plants. Yekarum, 9(2), 128\u2013146.",

    "Triebe, O., Hewamalage, H., Pilyugina, P., Laptev, N., Bergmeir, C., & Rajagopal, R. (2021). "
    "NeuralProphet: Explainable Forecasting at Scale. arXiv:2111.15397.",

    "Vargas-Forero, V. M., Manotas-Duque, D. F., & Trujillo, L. (2025). Comparative study of forecasting "
    "methods to predict the energy demand for the market of Colombia. International Journal of Energy "
    "Economics and Policy, 15(1), 65\u201376.",

    "Xie, J., Li, Z., Zhou, Z., & Liu, S. (2022). A hybrid forecasting model using LSTM and Prophet "
    "for energy consumption with decomposition of time series data. PeerJ Computer Science, 8, e1001.",
]

for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    r = p.add_run(f"[{i}] {ref}")
    r.font.size = Pt(9)
    r.font.name = 'Calibri'

# ============================================================
# SAVE
# ============================================================
doc.save(OUTPUT_PATH)
print(f"Literature Review saved to: {OUTPUT_PATH}")
size_kb = os.path.getsize(OUTPUT_PATH) / 1024
print(f"File size: {size_kb:.0f} KB")
