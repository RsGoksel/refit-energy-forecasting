"""
Generate concise 2-page Project Proposal (3 proposals, half page each).
Format: Title, Problem Statement, Proposed ML Approach, Expected Outcome.
"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE_DIR = os.path.dirname(__file__)
OUTPUT_PATH = os.path.join(BASE_DIR, "EBT629E_Project_Proposal.docx")

doc = Document()

# --- Page setup ---
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# --- Default style: Calibri 11pt, 1.3 line spacing ---
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.3

# ============================================================
# HEADER
# ============================================================
header = doc.add_paragraph()
header.alignment = WD_ALIGN_PARAGRAPH.CENTER
header.paragraph_format.space_after = Pt(2)
run = header.add_run("EBT 629E \u2013 Artificial Intelligence\n")
run.bold = True
run.font.size = Pt(13)
run.font.name = 'Calibri'

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_after = Pt(2)
run = sub.add_run("ML Applications in Energy Systems \u2013 Project Proposals")
run.bold = True
run.font.size = Pt(12)
run.font.name = 'Calibri'
run.font.color.rgb = RGBColor(0, 51, 102)

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.paragraph_format.space_after = Pt(8)
run = info.add_run(
    "Nurullah Y\u0131ld\u0131r\u0131m (301252004) | "
    "Kadir G\u00f6ksel G\u00fcnd\u00fcz (301241077) | "
    "Furkan \u00c7\u0131nar (301212001)\n"
    "April 6, 2026"
)
run.font.size = Pt(10)
run.font.name = 'Calibri'
run.font.color.rgb = RGBColor(80, 80, 80)

# --- Separator ---
sep = doc.add_paragraph()
sep.paragraph_format.space_after = Pt(6)
sep.paragraph_format.space_before = Pt(0)
run = sep.add_run("\u2500" * 85)
run.font.size = Pt(6)
run.font.color.rgb = RGBColor(180, 180, 180)


def add_project_title(number, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"Project {number}: {title}")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(0, 51, 102)


def add_section(heading, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(f"{heading}: ")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    run2 = p.add_run(text)
    run2.font.size = Pt(11)
    run2.font.name = 'Calibri'


def add_separator():
    sep = doc.add_paragraph()
    sep.paragraph_format.space_after = Pt(4)
    sep.paragraph_format.space_before = Pt(4)
    run = sep.add_run("\u2500" * 85)
    run.font.size = Pt(6)
    run.font.color.rgb = RGBColor(180, 180, 180)


# ============================================================
# PROJECT 1
# ============================================================
add_project_title(1, "Household Energy Consumption Forecasting Using Time Series Models")

add_section("Problem Statement",
    "Accurate prediction of household-level energy consumption is essential for effective demand-side "
    "management, load balancing, and integration of distributed energy resources in smart grids. "
    "However, residential consumption is inherently noisy and highly variable, driven by occupant "
    "behavior, appliance usage patterns, and temporal factors. Forecasting at this granularity remains "
    "challenging, and selecting the right modeling approach for different data characteristics is an "
    "open research question."
)

add_section("Proposed ML Approach",
    "We plan to use the REFIT Smart Home dataset, which contains high-resolution electrical load "
    "measurements (8-second sampling) from 20 UK households over two years (2013\u20132015), including "
    "aggregate and appliance-level readings. The raw data will be resampled to daily averages and "
    "smoothed using a rolling window to reduce noise. Three time series forecasting models will be "
    "compared: (1) ARIMA with automated parameter selection via ADF stationarity testing and AIC-based "
    "grid search, using a rolling one-step-ahead forecast strategy; (2) a multi-layer LSTM network "
    "with lookback windows to capture long-term temporal dependencies; and (3) FB-Prophet with weekly "
    "and yearly seasonality components. Models will be evaluated on an 80/20 temporal train-test split "
    "using MSE, RMSE, MAE, and R2 metrics."
)

add_section("Expected Outcome",
    "We expect ARIMA to perform well for short-term forecasting due to the strong autocorrelation "
    "in daily energy data, while LSTM may offer advantages when sufficient training data is available "
    "to learn complex nonlinear patterns. Prophet is expected to capture seasonal trends effectively. "
    "The comparison will provide practical guidance on model selection for household energy forecasting "
    "under varying data availability and noise conditions, contributing to optimized residential energy "
    "management strategies."
)

add_separator()

# ============================================================
# PROJECT 2
# ============================================================
add_project_title(2, "Non-Intrusive Load Monitoring (NILM) Using Deep Learning")

add_section("Problem Statement",
    "Non-Intrusive Load Monitoring (NILM) aims to disaggregate a household\u2019s total power signal "
    "into individual appliance-level consumption using only a single smart meter, eliminating the need "
    "for costly per-device sensors. Accurately separating overlapping and variable appliance signatures "
    "from a noisy aggregate signal remains a significant challenge, particularly for low-power and "
    "infrequently used devices."
)

add_section("Proposed ML Approach",
    "Using the REFIT dataset (which provides both aggregate and 9 individual appliance measurements as "
    "ground truth), we plan to train a sequence-to-point deep learning model combining 1D-CNN for "
    "local feature extraction and BiLSTM for temporal context. The aggregate signal will serve as input, "
    "and individual appliance power draws as targets. Transfer learning across multiple houses will be "
    "explored to improve generalization. Performance will be evaluated using F1-score for appliance "
    "state detection and MAE for energy estimation accuracy."
)

add_section("Expected Outcome",
    "A trained NILM model capable of disaggregating at least 5 major appliances (fridge, washing machine, "
    "dishwasher, kettle, microwave) with >80% F1-score, demonstrating the feasibility of software-based "
    "appliance monitoring as a scalable alternative to hardware-intensive solutions for residential energy "
    "management."
)

add_separator()

# ============================================================
# PROJECT 3
# ============================================================
add_project_title(3, "Short-Term Solar Irradiance Forecasting Using Hybrid CNN-LSTM")

add_section("Problem Statement",
    "Solar energy production depends directly on solar irradiance, which fluctuates with cloud cover, "
    "atmospheric conditions, and seasonal cycles. Accurate short-term forecasting (1\u20136 hours ahead) "
    "is critical for grid operators to manage PV intermittency and maintain supply-demand balance. "
    "Conventional numerical weather prediction models lack the spatial resolution and update frequency "
    "needed for reliable intra-day solar forecasting."
)

add_section("Proposed ML Approach",
    "We propose a hybrid CNN-LSTM architecture: convolutional layers will extract spatial features from "
    "geostationary satellite cloud imagery to capture cloud movement patterns, while LSTM layers will "
    "model the temporal evolution of irradiance time series. The model will be trained on publicly "
    "available solar irradiance datasets paired with satellite imagery. Attention mechanisms and data "
    "augmentation will be explored to improve accuracy under rapidly changing conditions. Evaluation "
    "metrics will include RMSE, MAE, and forecast skill score relative to persistence baselines."
)

add_section("Expected Outcome",
    "A forecasting system achieving at least 20% RMSE improvement over persistence baselines for 1\u20134 "
    "hour ahead predictions. By combining spatial (cloud morphology) and temporal (irradiance trends) "
    "information, the model is expected to outperform pure time series methods. This has direct "
    "applicability to Turkey\u2019s expanding solar energy sector for PV plant optimization and grid "
    "integration planning."
)

# ============================================================
# SAVE
# ============================================================
doc.save(OUTPUT_PATH)
print(f"Proposal saved to: {OUTPUT_PATH}")
size_kb = os.path.getsize(OUTPUT_PATH) / 1024
print(f"File size: {size_kb:.0f} KB")
